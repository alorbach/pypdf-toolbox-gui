"""
PDF Image Print Tool

Prepare images for double-sided printing by arranging them in front/back pairs.
Export to PDF or Word, or print directly. Supports mirroring, auto-trim,
and scaling to A4 format.

Copyright 2025-2026 Andre Lorbach

Licensed under the Apache License, Version 2.0 (the "License");
you may not use this file except in compliance with the License.
You may obtain a copy of the License at

    http://www.apache.org/licenses/LICENSE-2.0

Unless required by applicable law or agreed to in writing, software
distributed under the License is distributed on an "AS IS" BASIS,
WITHOUT WARRANTIES OR CONDITIONS OF ANY KIND, either express or implied.
See the License for the specific language governing permissions and
limitations under the License.
"""

import sys
import io
import re
import os
import tempfile
import subprocess
from pathlib import Path

import tkinter as tk
from tkinter import ttk, filedialog, messagebox, scrolledtext

from PIL import Image, ImageTk, ImageChops
from reportlab.lib.pagesizes import A4, landscape
from reportlab.lib.units import cm
from reportlab.pdfgen import canvas
from reportlab.lib.utils import ImageReader
from docx import Document
from docx.shared import Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT

try:
    from tkinterdnd2 import DND_FILES, TkinterDnD
    DND_AVAILABLE = True
except ImportError:
    DND_AVAILABLE = False

# ============================================================================
# Modern UI Styling Constants
# ============================================================================

class UIColors:
    """Modern color palette for consistent UI styling."""
    PRIMARY = "#2563eb"
    PRIMARY_HOVER = "#1d4ed8"
    PRIMARY_LIGHT = "#dbeafe"
    SECONDARY = "#64748b"
    SECONDARY_HOVER = "#475569"
    SUCCESS = "#16a34a"
    SUCCESS_LIGHT = "#dcfce7"
    SUCCESS_HOVER = "#15803d"
    ERROR = "#dc2626"
    ERROR_LIGHT = "#fee2e2"
    ERROR_HOVER = "#b91c1c"
    WARNING = "#f59e0b"
    WARNING_LIGHT = "#fef3c7"
    BG_PRIMARY = "#ffffff"
    BG_SECONDARY = "#f8fafc"
    BG_TERTIARY = "#f1f5f9"
    BORDER = "#e2e8f0"
    BORDER_DARK = "#cbd5e1"
    TEXT_PRIMARY = "#1e293b"
    TEXT_SECONDARY = "#64748b"
    TEXT_MUTED = "#94a3b8"
    DROP_ZONE_BG = "#f8fafc"
    DROP_ZONE_BORDER = "#94a3b8"
    DROP_ZONE_ACTIVE = "#dbeafe"
    DROP_ZONE_BORDER_ACTIVE = "#2563eb"
    THUMBNAIL_BG = "#ffffff"
    THUMBNAIL_HOVER = "#dbeafe"


class UIFonts:
    """Font configurations for consistent typography."""
    TITLE = ("Segoe UI", 18, "bold")
    SUBTITLE = ("Segoe UI", 14, "bold")
    HEADING = ("Segoe UI", 12, "bold")
    BODY = ("Segoe UI", 10)
    BODY_BOLD = ("Segoe UI", 10, "bold")
    SMALL = ("Segoe UI", 9)
    SMALL_BOLD = ("Segoe UI", 9, "bold")
    BUTTON = ("Segoe UI", 10, "bold")
    BUTTON_SMALL = ("Segoe UI", 9)


class UISpacing:
    """Consistent spacing values."""
    XS = 2
    SM = 5
    MD = 10
    LG = 15
    XL = 20
    XXL = 30


class PDFImagePrintTool:
    """Tool for preparing images for double-sided printing - export to PDF or Word."""

    def __init__(self):
        if DND_AVAILABLE:
            self.root = TkinterDnD.Tk()
        else:
            self.root = tk.Tk()

        self.root.title("PDF Image Print - Double-Sided Image Layout")
        self.root.minsize(800, 500)
        self.root.resizable(True, True)

        self.position_window()

        self.images = []
        self.image_mirrors = {}
        self.current_pair_index = 0
        self.debug_mode = tk.BooleanVar(value=False)
        self.margin = tk.DoubleVar(value=1.0)
        self.scale_to_width = tk.BooleanVar(value=True)
        self.auto_open_export = tk.BooleanVar(value=True)
        self.pdf_landscape = tk.BooleanVar(value=False)
        self.auto_trim = tk.BooleanVar(value=True)
        self.target_width = 29.7

        self.drag_start_index = None
        self.drag_start_y = None
        self.dragging = False
        self.drag_threshold = 5
        self.image_drag_source = None
        self.image_dragging = False
        self.preview_drag_source = None  # 'front' or 'back' when dragging from main preview
        self.debug_text = None

        self.setup_ui()
        self.log_debug("Application started")

    def position_window(self):
        """Position window in the area below the launcher."""
        x = int(os.environ.get('TOOL_WINDOW_X', 100))
        y = int(os.environ.get('TOOL_WINDOW_Y', 100))
        width = int(os.environ.get('TOOL_WINDOW_WIDTH', 1200))
        height = int(os.environ.get('TOOL_WINDOW_HEIGHT', 800))
        self.root.geometry(f"{width}x{height}+{x}+{y}")

    def setup_ui(self):
        main_frame = ttk.Frame(self.root, padding=str(UISpacing.MD))
        main_frame.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S))
        self.root.columnconfigure(0, weight=1)
        self.root.rowconfigure(0, weight=1)

        left_frame = ttk.LabelFrame(main_frame, text="Add images", padding=str(UISpacing.MD))
        left_frame.grid(row=0, column=0, rowspan=2, sticky=(tk.W, tk.E, tk.N, tk.S), padx=(0, 5))

        self.drop_area = tk.Text(
            left_frame, width=20, height=12, relief=tk.SUNKEN,
            bg=UIColors.DROP_ZONE_BG, fg=UIColors.TEXT_PRIMARY,
            font=UIFonts.BODY, wrap=tk.WORD,
            highlightbackground=UIColors.DROP_ZONE_BORDER,
            highlightthickness=2
        )
        self.drop_area.pack(fill=tk.BOTH, expand=True)
        self.drop_area.insert("1.0", "Drag images here\n\nOr click 'Select images'")
        self.drop_area.config(state=tk.DISABLED)

        if DND_AVAILABLE:
            self.drop_area.drop_target_register(DND_FILES)
            self.drop_area.dnd_bind('<<Drop>>', self.on_drop)
            self.drop_area.dnd_bind('<<DragEnter>>', self._on_drag_enter)
            self.drop_area.dnd_bind('<<DragLeave>>', self._on_drag_leave)

        ttk.Button(left_frame, text="Select images", command=self.select_images).pack(pady=UISpacing.SM)
        ttk.Button(left_frame, text="Clear all", command=self.clear_all).pack(pady=UISpacing.SM)

        middle_frame = ttk.LabelFrame(main_frame, text="Preview", padding=str(UISpacing.SM))
        middle_frame.grid(row=0, column=1, sticky=(tk.W, tk.E, tk.N, tk.S), padx=5)
        middle_frame.grid_rowconfigure(0, weight=1)
        middle_frame.grid_columnconfigure(0, weight=1)

        preview_canvas = tk.Canvas(middle_frame, bg=UIColors.BG_SECONDARY, highlightthickness=0)
        preview_scrollbar = ttk.Scrollbar(middle_frame, orient="vertical", command=preview_canvas.yview)
        preview_inner = ttk.Frame(preview_canvas)
        preview_inner.bind("<Configure>",
                          lambda e: preview_canvas.configure(scrollregion=preview_canvas.bbox("all")))
        preview_canvas_window = preview_canvas.create_window((0, 0), window=preview_inner, anchor="nw", tags="preview_inner")
        preview_canvas.configure(yscrollcommand=preview_scrollbar.set)
        preview_canvas.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S))
        preview_scrollbar.grid(row=0, column=1, sticky=(tk.N, tk.S))

        def _on_mousewheel_preview(event):
            preview_canvas.yview_scroll(int(-1 * (event.delta / 120)), "units")

        preview_canvas.bind("<MouseWheel>", _on_mousewheel_preview)
        preview_inner.bind("<MouseWheel>", _on_mousewheel_preview)

        def _preview_canvas_configure(event):
            if event.widget == preview_canvas and event.width > 1:
                preview_canvas.itemconfig(preview_canvas_window, width=event.width)

        preview_canvas.bind("<Configure>", _preview_canvas_configure)

        ttk.Label(preview_inner, text="Front:").pack(anchor=tk.W)
        front_frame = tk.Frame(preview_inner, bg=UIColors.BG_PRIMARY, relief=tk.SUNKEN, borderwidth=2,
                              highlightbackground=UIColors.BORDER, highlightthickness=1)
        front_frame.pack(pady=UISpacing.SM, fill=tk.BOTH, expand=True)
        self.front_preview = tk.Label(front_frame, bg=UIColors.BG_PRIMARY, cursor="hand2")
        self.front_preview.pack(expand=True, fill=tk.BOTH, padx=UISpacing.XS, pady=UISpacing.XS)

        ttk.Label(preview_inner, text="Back:").pack(anchor=tk.W)
        self.back_frame = tk.Frame(preview_inner, bg=UIColors.BG_PRIMARY, relief=tk.SUNKEN, borderwidth=2,
                             highlightbackground=UIColors.BORDER, highlightthickness=1)
        self.back_frame.pack(pady=UISpacing.SM, fill=tk.BOTH, expand=True)
        self.back_preview = tk.Label(self.back_frame, bg=UIColors.BG_PRIMARY, cursor="hand2")
        self.back_preview.pack(expand=True, fill=tk.BOTH, padx=UISpacing.XS, pady=UISpacing.XS)

        self.front_preview_frame = front_frame

        nav_frame = ttk.Frame(preview_inner)
        nav_frame.pack(pady=UISpacing.MD)
        ttk.Button(nav_frame, text="< Previous", command=self.prev_pair).pack(side=tk.LEFT, padx=UISpacing.SM)
        self.pair_label = ttk.Label(nav_frame, text="Pair 0 of 0")
        self.pair_label.pack(side=tk.LEFT, padx=UISpacing.SM)
        ttk.Button(nav_frame, text="Next >", command=self.next_pair).pack(side=tk.LEFT, padx=UISpacing.SM)

        ttk.Button(preview_inner, text="↔ Swap front/back",
                   command=lambda: self.swap_pair_images(self.current_pair_index)).pack(pady=UISpacing.SM)

        hint_label = ttk.Label(preview_inner, text="Drag Front ↔ Back to swap", font=UIFonts.SMALL,
                               foreground=UIColors.TEXT_MUTED)
        hint_label.pack(pady=(0, UISpacing.SM))

        self.preview_canvas = preview_canvas
        self.preview_inner = preview_inner

        self._setup_preview_drag()

        right_frame = ttk.LabelFrame(main_frame, text="Tiles – drag images between pairs to swap",
                                    padding=str(UISpacing.MD))
        right_frame.grid(row=0, column=2, sticky=(tk.W, tk.E, tk.N, tk.S), padx=(5, 0))

        canvas_frame = ttk.Frame(right_frame)
        canvas_frame.pack(fill=tk.BOTH, expand=True)
        self.tile_canvas = tk.Canvas(canvas_frame, bg=UIColors.BG_PRIMARY,
                                     highlightthickness=0)
        scrollbar = ttk.Scrollbar(canvas_frame, orient="vertical", command=self.tile_canvas.yview)
        self.tile_scrollable = ttk.Frame(self.tile_canvas)
        self.tile_scrollable.bind("<Configure>",
                                  lambda e: self.tile_canvas.configure(scrollregion=self.tile_canvas.bbox("all")))
        self.tile_canvas_window = self.tile_canvas.create_window((0, 0), window=self.tile_scrollable, anchor="nw")
        self.tile_canvas.configure(yscrollcommand=scrollbar.set)

        def _tile_canvas_configure(event):
            if event.widget == self.tile_canvas and event.width > 1:
                self.tile_canvas.itemconfig(self.tile_canvas_window, width=event.width)
        self.tile_canvas.bind("<Configure>", _tile_canvas_configure)
        self.tile_canvas.pack(side="left", fill="both", expand=True)
        scrollbar.pack(side="right", fill="y")

        settings_frame = ttk.LabelFrame(main_frame, text="Settings", padding=str(UISpacing.MD))
        settings_frame.grid(row=1, column=1, columnspan=2, sticky=(tk.W, tk.E), padx=5, pady=(5, 0))

        margin_frame = ttk.Frame(settings_frame)
        margin_frame.pack(fill=tk.X, pady=UISpacing.SM)
        ttk.Label(margin_frame, text="Margins (cm):").pack(side=tk.LEFT, padx=UISpacing.SM)
        margin_spin = ttk.Spinbox(margin_frame, from_=0.0, to=5.0, increment=0.1,
                                 textvariable=self.margin, width=10,
                                 command=lambda: self.update_previews())
        margin_spin.pack(side=tk.LEFT, padx=UISpacing.SM)
        self.margin.trace_add("write", lambda *a: self.update_previews())

        ttk.Checkbutton(settings_frame, text="Scale to A4 width (29.7 cm)",
                       variable=self.scale_to_width, command=self.update_previews).pack(anchor=tk.W, pady=UISpacing.SM)
        ttk.Checkbutton(settings_frame, text="Auto trim white borders",
                        variable=self.auto_trim).pack(anchor=tk.W, pady=UISpacing.SM)
        ttk.Checkbutton(settings_frame, text="PDF landscape (A4)",
                        variable=self.pdf_landscape).pack(anchor=tk.W, pady=UISpacing.SM)
        ttk.Checkbutton(settings_frame, text="Enable debug output",
                        variable=self.debug_mode, command=self.toggle_debug).pack(anchor=tk.W, pady=UISpacing.SM)
        ttk.Checkbutton(settings_frame, text="Auto open exported files",
                        variable=self.auto_open_export).pack(anchor=tk.W, pady=UISpacing.SM)

        action_frame = ttk.LabelFrame(main_frame, text="Actions", padding=str(UISpacing.MD))
        action_frame.grid(row=2, column=0, columnspan=3, sticky=(tk.W, tk.E), pady=(5, 0))

        ttk.Button(action_frame, text="Print", command=self.print_images).pack(side=tk.LEFT, padx=UISpacing.SM)
        ttk.Button(action_frame, text="Save as PDF", command=self.export_pdf).pack(side=tk.LEFT, padx=UISpacing.SM)
        ttk.Button(action_frame, text="Save as Word", command=self.export_word).pack(side=tk.LEFT, padx=UISpacing.SM)

        self.debug_frame = ttk.LabelFrame(main_frame, text="Debug output", padding=str(UISpacing.MD))
        self.debug_text = scrolledtext.ScrolledText(self.debug_frame, height=8, width=100, font=("Consolas", 9))
        self.debug_text.pack(fill=tk.BOTH, expand=True)
        self.debug_frame.grid_remove()

        main_frame.columnconfigure(0, minsize=160, weight=0)
        main_frame.columnconfigure(1, minsize=180, weight=1)
        main_frame.columnconfigure(2, minsize=180, weight=1)
        main_frame.rowconfigure(0, minsize=250, weight=1)
        main_frame.rowconfigure(1, minsize=0, weight=0)
        main_frame.rowconfigure(2, minsize=0, weight=0)

    def _on_drag_enter(self, event):
        self.drop_area.config(bg=UIColors.DROP_ZONE_ACTIVE,
                             highlightbackground=UIColors.DROP_ZONE_BORDER_ACTIVE)

    def _on_drag_leave(self, event):
        self.drop_area.config(bg=UIColors.DROP_ZONE_BG,
                             highlightbackground=UIColors.DROP_ZONE_BORDER)

    def _setup_preview_drag(self):
        """Enable drag-from-preview to swap front/back images."""
        def on_preview_press(side):
            def handler(e):
                self.preview_drag_source = side
                self.preview_drag_start_x = e.x_root
                self.preview_drag_start_y = e.y_root
            return handler

        def on_preview_release(source_side):
            def handler(e):
                if self.preview_drag_source is None:
                    return
                target_widget = self.back_preview if source_side == 'front' else self.front_preview
                if self._is_over_widget(e.x_root, e.y_root, target_widget):
                    if self.images and 0 <= self.current_pair_index < len(self.images):
                        moved = abs(e.x_root - self.preview_drag_start_x) + abs(e.y_root - self.preview_drag_start_y)
                        if moved > self.drag_threshold:
                            self.swap_pair_images(self.current_pair_index)
                self.preview_drag_source = None
            return handler

        self.front_preview.bind("<Button-1>", on_preview_press('front'))
        self.front_preview.bind("<ButtonRelease-1>", on_preview_release('front'))
        self.back_preview.bind("<Button-1>", on_preview_press('back'))
        self.back_preview.bind("<ButtonRelease-1>", on_preview_release('back'))

    def _is_over_widget(self, x_root, y_root, widget):
        """Check if screen coordinates are within the widget."""
        try:
            wx, wy = widget.winfo_rootx(), widget.winfo_rooty()
            ww, wh = widget.winfo_width(), widget.winfo_height()
            return wx <= x_root <= wx + ww and wy <= y_root <= wy + wh
        except Exception:
            return False

    def parse_dropped_files(self, data):
        """Parse dropped file paths (handles Windows {} wrapping)."""
        files = []
        if '{' in data:
            files = re.findall(r'\{([^}]+)\}', data)
            remaining = re.sub(r'\{[^}]+\}', '', data).strip()
            if remaining:
                files.extend(remaining.split())
        else:
            files = data.split()
        return [f for f in files if f.lower().endswith(('.png', '.jpg', '.jpeg', '.gif', '.bmp'))]

    def on_drop(self, event):
        files = self.parse_dropped_files(event.data)
        self._on_drag_leave(None)
        self.log_debug(f"Files received via drag & drop: {len(files)}")
        self.process_images(files)

    def log_debug(self, message):
        if self.debug_mode.get() and self.debug_text:
            self.debug_text.insert(tk.END, f"[DEBUG] {message}\n")
            self.debug_text.see(tk.END)
        if self.debug_mode.get():
            print(f"[DEBUG] {message}")

    def toggle_debug(self):
        if self.debug_mode.get():
            self.debug_frame.grid(row=3, column=0, columnspan=3, sticky=(tk.W, tk.E), pady=(5, 0))
        else:
            self.debug_frame.grid_remove()

    def select_images(self):
        files = filedialog.askopenfilenames(
            title="Select images",
            filetypes=[("Image files", "*.png *.jpg *.jpeg *.gif *.bmp"), ("All files", "*.*")]
        )
        if files:
            self.log_debug(f"Files selected: {len(files)}")
            self.process_images(list(files))

    def process_images(self, image_files):
        if not image_files:
            return
        for i in range(0, len(image_files), 2):
            front = image_files[i]
            back = image_files[i + 1] if i + 1 < len(image_files) else None
            self.images.append((front, back))
            self.log_debug(f"Pair added: front={front}, back={back}")
        self.update_previews()
        self.update_tile_view()
        self.log_debug(f"Total pairs: {len(self.images)}")

    def clear_all(self):
        self.images = []
        self.image_mirrors = {}
        self.current_pair_index = 0
        self.update_previews()
        self.update_tile_view()
        self.log_debug("All images cleared")

    def update_previews(self):
        if not self.images:
            self.front_preview.config(image='')
            self.back_preview.config(image='')
            if self.pair_label:
                self.pair_label.config(text="Pair 0 of 0")
            return
        if 0 <= self.current_pair_index < len(self.images):
            front_path, back_path = self.images[self.current_pair_index]
            if front_path:
                self.show_preview(front_path, self.front_preview,
                                 pair_index=self.current_pair_index, side='front')
                self.front_preview.bind("<Button-3>", lambda e: self.show_image_menu(e, self.current_pair_index, 'front'))
            else:
                self.front_preview.config(image='')
                self.front_preview.unbind("<Button-3>")
            if back_path:
                self.show_preview(back_path, self.back_preview,
                                 pair_index=self.current_pair_index, side='back')
                self.back_preview.bind("<Button-3>", lambda e: self.show_image_menu(e, self.current_pair_index, 'back'))
            else:
                self.back_preview.config(image='')
                self.back_preview.unbind("<Button-3>")
            if self.pair_label:
                self.pair_label.config(text=f"Pair {self.current_pair_index + 1} of {len(self.images)}")

    def show_preview(self, image_path, label_widget, max_size=(400, 300), pair_index=None, side=None):
        try:
            img = Image.open(image_path)
            if pair_index is not None and side is not None:
                mirror = self.image_mirrors.get((pair_index, side), 'none')
                img = self.apply_mirror(img, mirror)
            img.thumbnail(max_size, Image.Resampling.LANCZOS)
            photo = ImageTk.PhotoImage(img)
            label_widget.config(image=photo)
            label_widget.image = photo
        except Exception as e:
            self.log_debug(f"Failed to load {image_path}: {e}")
            messagebox.showerror("Error", f"Could not load image: {e}")

    def apply_mirror(self, img, mirror_type):
        if mirror_type == 'h' or mirror_type == 'horizontal':
            return img.transpose(Image.Transpose.FLIP_LEFT_RIGHT)
        elif mirror_type == 'v' or mirror_type == 'vertical':
            return img.transpose(Image.Transpose.FLIP_TOP_BOTTOM)
        elif mirror_type == 'both':
            img = img.transpose(Image.Transpose.FLIP_LEFT_RIGHT)
            return img.transpose(Image.Transpose.FLIP_TOP_BOTTOM)
        return img

    def trim_image(self, img):
        try:
            if img.mode in ("RGBA", "LA"):
                alpha = img.split()[-1]
                bbox = alpha.getbbox()
                if bbox:
                    return img.crop(bbox)
            rgb = img.convert("RGB")
            bg = Image.new("RGB", rgb.size, (255, 255, 255))
            diff = ImageChops.difference(rgb, bg)
            bbox = diff.getbbox()
            if bbox:
                return img.crop(bbox)
        except Exception as e:
            self.log_debug(f"Trim failed: {e}")
        return img

    def load_base_image(self, image_path, pair_index=None, side=None, mirror=False, trim=False):
        img = Image.open(image_path)
        if pair_index is not None and side is not None:
            mirror_type = self.image_mirrors.get((pair_index, side), 'none')
            img = self.apply_mirror(img, mirror_type)
        elif mirror:
            img = img.transpose(Image.Transpose.FLIP_LEFT_RIGHT)
        if trim:
            img = self.trim_image(img)
        return img

    def compute_target_size_cm(self, img, available_width_cm, available_height_cm):
        aspect_ratio = img.height / img.width
        width_cm = available_width_cm
        height_cm = width_cm * aspect_ratio
        if height_cm > available_height_cm:
            height_cm = available_height_cm
            width_cm = height_cm / aspect_ratio
        return width_cm, height_cm

    def prepare_export_image(self, img, target_width_cm=None, target_height_cm=None):
        if target_width_cm is None or target_height_cm is None:
            return img
        target_dpi = 300
        target_width_px = int(target_width_cm / 2.54 * target_dpi)
        target_height_px = int(target_height_cm / 2.54 * target_dpi)
        return img.resize((target_width_px, target_height_px), Image.Resampling.LANCZOS)

    def update_tile_view(self):
        for widget in self.tile_scrollable.winfo_children():
            widget.destroy()
        for idx, (front_path, back_path) in enumerate(self.images):
            pair_frame = ttk.Frame(self.tile_scrollable, relief=tk.RAISED, borderwidth=2)
            pair_frame.pack(fill=tk.X, padx=UISpacing.SM, pady=UISpacing.SM)
            pair_frame.pair_index = idx

            drag_handle_frame = tk.Frame(pair_frame, height=20, bg=UIColors.PRIMARY_LIGHT, cursor="hand2")
            drag_handle_frame.pack(fill=tk.X)
            drag_handle_frame.pair_index = idx
            drag_label = tk.Label(drag_handle_frame, text="☰ Drag to reorder",
                                 bg=UIColors.PRIMARY_LIGHT, fg=UIColors.PRIMARY, font=UIFonts.SMALL)
            drag_label.pack()
            drag_label.pair_index = idx

            content_frame = ttk.Frame(pair_frame)
            content_frame.pack(fill=tk.BOTH, expand=True)
            title_frame = ttk.Frame(content_frame)
            title_frame.pack()
            ttk.Label(title_frame, text=f"Pair {idx + 1}", font=UIFonts.BODY_BOLD).pack(side=tk.LEFT, padx=UISpacing.SM)
            ttk.Button(title_frame, text="↔ Swap", width=10,
                      command=lambda i=idx: self.swap_pair_images(i)).pack(side=tk.LEFT, padx=UISpacing.SM)

            front_frame = ttk.Frame(content_frame)
            front_frame.pack(side=tk.LEFT, padx=UISpacing.SM, pady=UISpacing.SM)
            ttk.Label(front_frame, text="Front").pack()
            front_tile_frame = tk.Frame(front_frame, bg=UIColors.BG_PRIMARY, relief=tk.SUNKEN, borderwidth=1)
            front_tile_frame.pack()
            front_tile = tk.Label(front_tile_frame, bg=UIColors.BG_PRIMARY,
                                 cursor="hand2" if front_path else "")
            front_tile.pack(padx=UISpacing.XS, pady=UISpacing.XS)
            if front_path:
                self.show_preview(front_path, front_tile, max_size=(120, 120), pair_index=idx, side='front')
                front_tile.bind("<Button-3>", lambda e, i=idx: self.show_image_menu(e, i, 'front'))

            back_frame = ttk.Frame(content_frame)
            back_frame.pack(side=tk.LEFT, padx=UISpacing.SM, pady=UISpacing.SM)
            ttk.Label(back_frame, text="Back").pack()
            back_tile_frame = tk.Frame(back_frame, bg=UIColors.BG_PRIMARY, relief=tk.SUNKEN, borderwidth=1)
            back_tile_frame.pack()
            back_tile = tk.Label(back_tile_frame, bg=UIColors.BG_PRIMARY,
                                cursor="hand2" if back_path else "")
            back_tile.pack(padx=UISpacing.XS, pady=UISpacing.XS)
            if back_path:
                self.show_preview(back_path, back_tile, max_size=(120, 120), pair_index=idx, side='back')
                back_tile.bind("<Button-3>", lambda e, i=idx: self.show_image_menu(e, i, 'back'))
            else:
                ttk.Label(back_tile_frame, text="(empty)", bg=UIColors.BG_PRIMARY).pack(padx=UISpacing.XS, pady=UISpacing.XS)

            def make_drag_start(i):
                return lambda e: self.on_drag_start(e, i)
            def make_drag_motion(i):
                return lambda e: self.on_drag_motion(e, i)
            def make_drag_end(i):
                return lambda e: self.on_drag_end(e, i)
            def make_click_handler(i):
                return lambda e: self.select_pair(i)
            def make_pair_menu_handler(i):
                return lambda e: self.show_pair_menu(e, i)

            drag_handle_frame.bind("<Button-1>", make_drag_start(idx))
            drag_handle_frame.bind("<B1-Motion>", make_drag_motion(idx))
            drag_handle_frame.bind("<ButtonRelease-1>", make_drag_end(idx))
            drag_label.bind("<Button-1>", make_drag_start(idx))
            drag_label.bind("<B1-Motion>", make_drag_motion(idx))
            drag_label.bind("<ButtonRelease-1>", make_drag_end(idx))
            content_frame.bind("<Button-1>", make_click_handler(idx))
            for child in content_frame.winfo_children():
                if isinstance(child, (ttk.Frame, tk.Frame, tk.Label, ttk.Label)):
                    child.bind("<Button-1>", make_click_handler(idx))
            pair_frame.bind("<Button-3>", make_pair_menu_handler(idx))
            content_frame.bind("<Button-3>", make_pair_menu_handler(idx))

            if front_path:
                front_tile.bind("<Button-1>", lambda e, i=idx, s='front': self.on_image_drag_start(e, i, s))
                front_tile.bind("<B1-Motion>", lambda e, i=idx, s='front': self.on_image_drag_motion(e, i, s))
                front_tile.bind("<ButtonRelease-1>", lambda e, i=idx, s='front': self.on_image_drag_end(e, i, s))
                front_tile_frame.bind("<Button-1>", lambda e, i=idx, s='front': self.on_image_drag_start(e, i, s))
                front_tile_frame.bind("<B1-Motion>", lambda e, i=idx, s='front': self.on_image_drag_motion(e, i, s))
                front_tile_frame.bind("<ButtonRelease-1>", lambda e, i=idx, s='front': self.on_image_drag_end(e, i, s))
            if back_path:
                back_tile.bind("<Button-1>", lambda e, i=idx, s='back': self.on_image_drag_start(e, i, s))
                back_tile.bind("<B1-Motion>", lambda e, i=idx, s='back': self.on_image_drag_motion(e, i, s))
                back_tile.bind("<ButtonRelease-1>", lambda e, i=idx, s='back': self.on_image_drag_end(e, i, s))
                back_tile_frame.bind("<Button-1>", lambda e, i=idx, s='back': self.on_image_drag_start(e, i, s))
                back_tile_frame.bind("<B1-Motion>", lambda e, i=idx, s='back': self.on_image_drag_motion(e, i, s))
                back_tile_frame.bind("<ButtonRelease-1>", lambda e, i=idx, s='back': self.on_image_drag_end(e, i, s))

    def select_pair(self, index):
        if not self.dragging:
            self.current_pair_index = index
            self.update_previews()
        self.log_debug(f"Pair {index + 1} selected")

    def on_drag_start(self, event, index):
        self.drag_start_index = index
        self.drag_start_y = event.y_root
        self.dragging = False

    def on_drag_motion(self, event, index):
        if self.drag_start_index is None:
            return
        if abs(event.y_root - self.drag_start_y) > self.drag_threshold:
            self.dragging = True
            target_index = self.find_drop_position(event.y_root)
            if target_index != self.drag_start_index:
                self.update_drag_feedback(self.drag_start_index, target_index)

    def on_drag_end(self, event, index):
        if self.drag_start_index is None:
            return
        if self.dragging:
            target_index = self.find_drop_position(event.y_root)
            if target_index != self.drag_start_index and target_index is not None:
                self.reorder_pairs(self.drag_start_index, target_index)
        self.drag_start_index = None
        self.drag_start_y = None
        self.dragging = False
        self.update_tile_view()
        self.update_previews()

    def find_drop_position(self, y_root):
        pair_frames = []
        for widget in self.tile_scrollable.winfo_children():
            if hasattr(widget, 'pair_index'):
                try:
                    wy = widget.winfo_rooty()
                    wh = widget.winfo_height()
                    pair_frames.append({
                        'index': widget.pair_index, 'y': wy, 'height': wh,
                        'center': wy + wh / 2, 'top': wy, 'bottom': wy + wh
                    })
                except Exception:
                    pair_frames.append({
                        'index': widget.pair_index, 'y': widget.pair_index * 200,
                        'height': 200, 'center': widget.pair_index * 200 + 100,
                        'top': widget.pair_index * 200, 'bottom': (widget.pair_index + 1) * 200
                    })
        if not pair_frames:
            return self.drag_start_index
        pair_frames.sort(key=lambda x: x['y'])
        for fi in pair_frames:
            if fi['top'] <= y_root <= fi['bottom']:
                if y_root < fi['center']:
                    pos = pair_frames.index(fi)
                    return pair_frames[pos - 1]['index'] if pos > 0 else fi['index']
                return fi['index']
        return pair_frames[0]['index'] if y_root < pair_frames[0]['top'] else pair_frames[-1]['index']

    def reorder_pairs(self, from_index, to_index):
        if from_index == to_index:
            return
        pair = self.images.pop(from_index)
        self.images.insert(to_index, pair)
        if self.current_pair_index == from_index:
            self.current_pair_index = to_index
        elif from_index < self.current_pair_index <= to_index:
            self.current_pair_index -= 1
        elif to_index <= self.current_pair_index < from_index:
            self.current_pair_index += 1

    def update_drag_feedback(self, from_index, to_index):
        for widget in self.tile_scrollable.winfo_children():
            if hasattr(widget, 'pair_index'):
                if widget.pair_index == from_index:
                    widget.config(relief=tk.SUNKEN, borderwidth=3)
                elif widget.pair_index == to_index:
                    widget.config(relief=tk.RIDGE, borderwidth=3)
                else:
                    widget.config(relief=tk.RAISED, borderwidth=2)

    def prev_pair(self):
        if self.images and self.current_pair_index > 0:
            self.current_pair_index -= 1
            self.update_previews()

    def next_pair(self):
        if self.images and self.current_pair_index < len(self.images) - 1:
            self.current_pair_index += 1
            self.update_previews()

    def show_image_menu(self, event, pair_index, side):
        menu = tk.Menu(self.root, tearoff=0)
        current = self.image_mirrors.get((pair_index, side), 'none')
        status = {"none": "No mirroring", "h": "Horizontal", "v": "Vertical", "both": "Both"}[current]
        menu.add_command(label=f"Current: {status}", state=tk.DISABLED)
        menu.add_separator()
        menu.add_command(label="No mirroring", command=lambda: self.set_image_mirror(pair_index, side, 'none'))
        menu.add_command(label="Mirror horizontally", command=lambda: self.set_image_mirror(pair_index, side, 'h'))
        menu.add_command(label="Mirror vertically", command=lambda: self.set_image_mirror(pair_index, side, 'v'))
        menu.add_command(label="Mirror both", command=lambda: self.set_image_mirror(pair_index, side, 'both'))
        try:
            menu.tk_popup(event.x_root, event.y_root)
        finally:
            menu.grab_release()

    def set_image_mirror(self, pair_index, side, mirror_type):
        if mirror_type == 'none':
            if (pair_index, side) in self.image_mirrors:
                del self.image_mirrors[(pair_index, side)]
        else:
            self.image_mirrors[(pair_index, side)] = mirror_type
        self.update_previews()
        self.update_tile_view()

    def swap_pair_images(self, pair_index):
        if 0 <= pair_index < len(self.images):
            front_path, back_path = self.images[pair_index]
            self.images[pair_index] = (back_path, front_path)
            fm = self.image_mirrors.pop((pair_index, 'front'), None)
            bm = self.image_mirrors.pop((pair_index, 'back'), None)
            if fm:
                self.image_mirrors[(pair_index, 'back')] = fm
            if bm:
                self.image_mirrors[(pair_index, 'front')] = bm
            if self.current_pair_index == pair_index:
                self.update_previews()
            self.update_tile_view()

    def show_pair_menu(self, event, pair_index):
        menu = tk.Menu(self.root, tearoff=0)
        menu.add_command(label=f"Delete pair {pair_index + 1}", command=lambda: self.delete_pair(pair_index))
        menu.add_separator()
        menu.add_command(label="Cancel")
        try:
            menu.tk_popup(event.x_root, event.y_root)
        finally:
            menu.grab_release()

    def delete_pair(self, pair_index):
        if 0 <= pair_index < len(self.images):
            for key in list(self.image_mirrors.keys()):
                if key[0] == pair_index:
                    del self.image_mirrors[key]
            self.images.pop(pair_index)
            new_mirrors = {}
            for (idx, side), mirror_type in self.image_mirrors.items():
                if idx > pair_index:
                    new_mirrors[(idx - 1, side)] = mirror_type
                elif idx < pair_index:
                    new_mirrors[(idx, side)] = mirror_type
            self.image_mirrors = new_mirrors
            if self.current_pair_index >= len(self.images):
                self.current_pair_index = max(0, len(self.images) - 1)
            elif self.current_pair_index > pair_index:
                self.current_pair_index -= 1
            self.update_previews()
            self.update_tile_view()

    def on_image_drag_start(self, event, pair_index, side):
        self.image_drag_source = (pair_index, side)
        self.image_drag_start_x = event.x_root
        self.image_drag_start_y = event.y_root
        self.image_dragging = False

    def on_image_drag_motion(self, event, pair_index, side):
        if self.image_drag_source is None:
            return
        if (abs(event.x_root - self.image_drag_start_x) > self.drag_threshold or
                abs(event.y_root - self.image_drag_start_y) > self.drag_threshold):
            self.image_dragging = True

    def on_image_drag_end(self, event, pair_index, side):
        if self.image_drag_source is None:
            return
        src_pair, src_side = self.image_drag_source
        tgt_pair, tgt_side = self.find_image_at_position(event.x_root, event.y_root)
        if tgt_pair is not None and tgt_side is not None and (src_pair, src_side) != (tgt_pair, tgt_side):
            self.swap_images_between_pairs(src_pair, src_side, tgt_pair, tgt_side)
        self.image_drag_source = None
        self.image_dragging = False

    def find_image_at_position(self, x_root, y_root):
        for widget in self.tile_scrollable.winfo_children():
            if hasattr(widget, 'pair_index'):
                try:
                    wx, wy = widget.winfo_rootx(), widget.winfo_rooty()
                    ww, wh = widget.winfo_width(), widget.winfo_height()
                    if wx <= x_root <= wx + ww and wy <= y_root <= wy + wh:
                        rx = x_root - wx
                        return (widget.pair_index, 'front' if rx < ww / 2 else 'back')
                except Exception:
                    pass
        return (None, None)

    def swap_images_between_pairs(self, source_pair, source_side, target_pair, target_side):
        if (source_pair == target_pair and source_side == target_side) or \
           source_pair >= len(self.images) or target_pair >= len(self.images):
            return
        sf, sb = self.images[source_pair]
        tf, tb = self.images[target_pair]
        si = sf if source_side == 'front' else sb
        ti = tf if target_side == 'front' else tb
        if source_side == 'front':
            if target_side == 'front':
                self.images[source_pair] = (ti, sb)
                self.images[target_pair] = (si, tb)
            else:
                self.images[source_pair] = (ti, sb)
                self.images[target_pair] = (tf, si)
        else:
            if target_side == 'front':
                self.images[source_pair] = (sf, ti)
                self.images[target_pair] = (si, tb)
            else:
                self.images[source_pair] = (sf, ti)
                self.images[target_pair] = (tf, si)
        sm = self.image_mirrors.pop((source_pair, source_side), None)
        tm = self.image_mirrors.pop((target_pair, target_side), None)
        if sm:
            self.image_mirrors[(target_pair, target_side)] = sm
        if tm:
            self.image_mirrors[(source_pair, source_side)] = tm
        if self.current_pair_index in [source_pair, target_pair]:
            self.update_previews()
        self.update_tile_view()

    def create_pdf(self, filename):
        page_size = landscape(A4) if self.pdf_landscape.get() else A4
        c = canvas.Canvas(filename, pagesize=page_size)
        width, height = page_size
        pw = 29.7 if self.pdf_landscape.get() else 21.0
        ph = 21.0 if self.pdf_landscape.get() else 29.7
        margin = self.margin.get() * cm
        for idx, (front_path, back_path) in enumerate(self.images):
            if front_path:
                base = self.load_base_image(front_path, pair_index=idx, side='front',
                                            mirror=False, trim=self.auto_trim.get())
                if self.scale_to_width.get():
                    mc = self.margin.get()
                    aw, ah = pw - 2 * mc, ph - 2 * mc
                    iw, ih = self.compute_target_size_cm(base, aw, ah)
                    exp = self.prepare_export_image(base, iw, ih)
                    img_w, img_h = iw * cm, ih * cm
                else:
                    exp = base
                    dpi = 72
                    img_w = (exp.width / dpi) * 2.54 * cm
                    img_h = (exp.height / dpi) * 2.54 * cm
                x, y = (width - img_w) / 2, (height - img_h) / 2
                c.drawImage(ImageReader(exp), x, y, width=img_w, height=img_h)
            c.showPage()
            if back_path:
                base = self.load_base_image(back_path, pair_index=idx, side='back',
                                            mirror=False, trim=self.auto_trim.get())
                if self.scale_to_width.get():
                    mc = self.margin.get()
                    aw, ah = pw - 2 * mc, ph - 2 * mc
                    iw, ih = self.compute_target_size_cm(base, aw, ah)
                    exp = self.prepare_export_image(base, iw, ih)
                    img_w, img_h = iw * cm, ih * cm
                else:
                    exp = base
                    dpi = 72
                    img_w = (exp.width / dpi) * 2.54 * cm
                    img_h = (exp.height / dpi) * 2.54 * cm
                x, y = (width - img_w) / 2, (height - img_h) / 2
                c.drawImage(ImageReader(exp), x, y, width=img_w, height=img_h)
            c.showPage()
        c.save()

    def open_file(self, filepath):
        try:
            if sys.platform == 'win32':
                try:
                    os.startfile(filepath)
                except Exception:
                    subprocess.Popen(["cmd", "/c", "start", "", filepath], shell=True)
            elif sys.platform == 'darwin':
                os.system(f'open "{filepath}"')
            else:
                os.system(f'xdg-open "{filepath}"')
            return True
        except Exception:
            return False

    def print_images(self):
        if not self.images:
            messagebox.showwarning("Warning", "No images to print.")
            return
        try:
            tmp = tempfile.NamedTemporaryFile(delete=False, suffix='.pdf')
            tmp.close()
            self.create_pdf(tmp.name)
            if sys.platform == 'win32':
                os.startfile(tmp.name)
            elif sys.platform == 'darwin':
                os.system(f'open -a "Preview" "{tmp.name}"')
            else:
                os.system(f'xdg-open "{tmp.name}"')
        except Exception as e:
            self.log_debug(f"Print failed: {e}")
            messagebox.showerror("Error", f"Print failed: {e}")

    def export_pdf(self):
        if not self.images:
            messagebox.showwarning("Warning", "No images to export.")
            return
        filename = filedialog.asksaveasfilename(
            defaultextension=".pdf",
            filetypes=[("PDF files", "*.pdf"), ("All files", "*.*")]
        )
        if filename:
            try:
                self.root.config(cursor="wait")
                self.root.update()
                self.create_pdf(filename)
                if self.auto_open_export.get():
                    self.open_file(filename)
            except Exception as e:
                self.log_debug(f"PDF export failed: {e}")
                messagebox.showerror("Error", f"Save failed: {e}")
            finally:
                self.root.config(cursor="")

    def export_word(self):
        if not self.images:
            messagebox.showwarning("Warning", "No images to export.")
            return
        filename = filedialog.asksaveasfilename(
            defaultextension=".docx",
            filetypes=[("Word documents", "*.docx"), ("All files", "*.*")]
        )
        if filename:
            try:
                self.root.config(cursor="wait")
                self.root.update()
                doc = Document()
                section = doc.sections[0]
                section.orientation = WD_ORIENT.LANDSCAPE
                section.page_width = Cm(29.7)
                section.page_height = Cm(21.0)
                section.left_margin = Cm(self.margin.get())
                section.right_margin = Cm(self.margin.get())
                section.top_margin = Cm(self.margin.get())
                section.bottom_margin = Cm(self.margin.get())
                aw = (section.page_width - section.left_margin - section.right_margin).cm
                ah = (section.page_height - section.top_margin - section.bottom_margin).cm
                for idx, (front_path, back_path) in enumerate(self.images):
                    if front_path:
                        base = self.load_base_image(front_path, pair_index=idx, side='front',
                                                    mirror=False, trim=self.auto_trim.get())
                        img_bytes = io.BytesIO()
                        base.save(img_bytes, format='PNG')
                        img_bytes.seek(0)
                        p = doc.add_paragraph()
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run = p.add_run()
                        if self.scale_to_width.get():
                            iw, ih = self.compute_target_size_cm(base, aw, ah)
                            run.add_picture(img_bytes, width=Cm(iw), height=Cm(ih))
                        else:
                            run.add_picture(img_bytes, width=Cm(aw))
                    doc.add_page_break()
                    if back_path:
                        base = self.load_base_image(back_path, pair_index=idx, side='back',
                                                    mirror=False, trim=self.auto_trim.get())
                        img_bytes = io.BytesIO()
                        base.save(img_bytes, format='PNG')
                        img_bytes.seek(0)
                        p = doc.add_paragraph()
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run = p.add_run()
                        if self.scale_to_width.get():
                            iw, ih = self.compute_target_size_cm(base, aw, ah)
                            run.add_picture(img_bytes, width=Cm(iw), height=Cm(ih))
                        else:
                            run.add_picture(img_bytes, width=Cm(aw))
                    else:
                        doc.add_paragraph().add_run("(No back side)")
                    if idx < len(self.images) - 1:
                        doc.add_page_break()
                doc.save(filename)
                if self.auto_open_export.get():
                    self.open_file(filename)
            except Exception as e:
                self.log_debug(f"Word export failed: {e}")
                messagebox.showerror("Error", f"Save failed: {e}")
            finally:
                self.root.config(cursor="")

    def run(self):
        self.root.mainloop()


def main():
    app = PDFImagePrintTool()
    app.run()


if __name__ == "__main__":
    main()
