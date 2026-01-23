"""
PDF Markdown Converter Tool

Convert Markdown files to PDF or DOCX documents with modern styling and formatting.

Copyright 2025-2026 Andre Lorbach

Licensed under the Apache License, Version 2.0 (the "License");
you may not use this file except in compliance with the License.
You may obtain a copy of the License at

    http://www.apache.org/licenses/LICENSE-2.0

Unless required by applicable law or agreed to in writing, software
distributed under the License is distributed on an "AS IS" BASIS,
WITHOUT WARRANTIES OR CONDITIONS OF ANY KIND, either express or implied.
See the License for the specific language governing permissions and
limitations under the License.
"""

import os
import sys
import re
import shutil
import subprocess
import tempfile
import unicodedata
import webbrowser
import threading
import time
import atexit
import json
import io
from pathlib import Path
from typing import Optional, Dict, Any, Tuple, List

import tkinter as tk
from tkinter import ttk, filedialog, messagebox, scrolledtext

# Drag and drop support
try:
    from tkinterdnd2 import DND_FILES, TkinterDnD
    HAS_DND = True
except ImportError:
    HAS_DND = False
    print("[WARNING] tkinterdnd2 not available, drag and drop disabled")

# Markdown processing
try:
    import markdown
    from markdown.extensions import codehilite, tables, toc, fenced_code
    MARKDOWN_AVAILABLE = True
except ImportError:
    MARKDOWN_AVAILABLE = False
    print("[WARNING] 'markdown' library not found. Install with: pip install markdown")
    print("[WARNING] Tool will start but markdown conversion will be disabled.")

# PDF generation with ReportLab
try:
    from reportlab.lib.pagesizes import letter, A4, landscape, portrait
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak, Preformatted
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib import colors
    from reportlab.lib.units import inch
    from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_JUSTIFY
    from reportlab.pdfgen import canvas
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    REPORTLAB_AVAILABLE = True
except ImportError:
    REPORTLAB_AVAILABLE = False
    print("[INFO] 'reportlab' not found. PDF export will be limited.")

# WeasyPrint (optional, lazy import)
weasyprint = None
WEASYPRINT_AVAILABLE = False

def _setup_weasyprint_dll_directories():
    """Setup GTK DLL directories for WeasyPrint on Windows (Feature 6)"""
    if sys.platform != 'win32':
        return
    
    # Check if WEASYPRINT_DLL_DIRECTORIES is already set
    dll_dirs = os.environ.get('WEASYPRINT_DLL_DIRECTORIES', '')
    
    # Check system environment variables (Windows)
    if not dll_dirs:
        try:
            import winreg
            # Check machine-level environment variables
            try:
                with winreg.OpenKey(winreg.HKEY_LOCAL_MACHINE, r'SYSTEM\CurrentControlSet\Control\Session Manager\Environment') as key:
                    dll_dirs = winreg.QueryValueEx(key, 'WEASYPRINT_DLL_DIRECTORIES')[0] or ''
            except (FileNotFoundError, OSError):
                pass
            
            # Check user-level environment variables
            if not dll_dirs:
                try:
                    with winreg.OpenKey(winreg.HKEY_CURRENT_USER, r'Environment') as key:
                        dll_dirs = winreg.QueryValueEx(key, 'WEASYPRINT_DLL_DIRECTORIES')[0] or ''
                except (FileNotFoundError, OSError):
                    pass
        except ImportError:
            pass  # winreg not available
    
    # If not set, try to find GTK in common locations
    if not dll_dirs:
        possible_gtk_paths = [
            r'C:\Program Files\GTK3-Runtime Win64\bin',
            r'C:\Program Files (x86)\GTK3-Runtime Win64\bin',
            r'C:\GTK3-Runtime\bin',
            r'C:\Program Files\GTK\bin',
            r'C:\Program Files (x86)\GTK\bin',
            r'C:\msys64\mingw64\bin',  # MSYS2/MinGW location
            r'C:\msys64\usr\bin',  # MSYS2 location
        ]
        
        # Also check PATH for GTK
        path_dirs = os.environ.get('PATH', '').split(os.pathsep)
        for path_dir in path_dirs:
            if path_dir and os.path.isdir(path_dir):
                try:
                    files = os.listdir(path_dir)
                    if any('gobject' in f.lower() and f.endswith('.dll') for f in files):
                        possible_gtk_paths.insert(0, path_dir)
                except (OSError, PermissionError):
                    pass
        
        # Find first existing GTK bin directory with gobject DLL
        for gtk_path in possible_gtk_paths:
            if os.path.isdir(gtk_path):
                try:
                    dll_files = [f for f in os.listdir(gtk_path) if 'gobject' in f.lower() and f.endswith('.dll')]
                    if dll_files:
                        dll_dirs = gtk_path
                        break
                except (OSError, PermissionError):
                    continue
    
    # Set WEASYPRINT_DLL_DIRECTORIES if we found a path
    if dll_dirs:
        os.environ['WEASYPRINT_DLL_DIRECTORIES'] = dll_dirs
        # Also add to DLL search path for current process
        if hasattr(os, 'add_dll_directory'):
            try:
                os.add_dll_directory(dll_dirs)
            except Exception:
                pass

def _try_import_weasyprint():
    """Safely try to import WeasyPrint only when needed (Feature 6)"""
    global weasyprint, WEASYPRINT_AVAILABLE
    
    if WEASYPRINT_AVAILABLE:
        return True
    
    # Setup DLL directories for Windows
    _setup_weasyprint_dll_directories()
    
    # Temporarily suppress stderr to catch WeasyPrint import errors silently
    old_stderr = sys.stderr
    sys.stderr = io.StringIO()
    try:
        from weasyprint import HTML, CSS
        from weasyprint.text.fonts import FontConfiguration
        weasyprint = {'HTML': HTML, 'CSS': CSS, 'FontConfiguration': FontConfiguration}
        WEASYPRINT_AVAILABLE = True
        sys.stderr = old_stderr
        return True
    except Exception:
        sys.stderr = old_stderr
        WEASYPRINT_AVAILABLE = False
        return False

# Embedded HTML browser widget support
try:
    from tkinterweb import HtmlFrame
    TKINTERWEB_AVAILABLE = True
except ImportError:
    TKINTERWEB_AVAILABLE = False
    print("[INFO] tkinterweb not available. Using text preview instead.")

# DOCX support
try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.style import WD_STYLE_TYPE
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import RGBColor
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("[WARNING] 'python-docx' not found. DOCX export disabled.")

# HTML to text conversion (optional)
try:
    import html2text
    HTML2TEXT_AVAILABLE = True
except ImportError:
    HTML2TEXT_AVAILABLE = False


# ============================================================================
# Encoding Detection
# ============================================================================

def _read_file_with_encoding_detection(file_path):
    """Read file with automatic encoding detection"""
    # First, check for BOM (Byte Order Mark) to detect UTF-16/UTF-8-BOM
    with open(file_path, 'rb') as f:
        first_bytes = f.read(4)
        
        # Check for UTF-16 BOMs
        if first_bytes.startswith(b'\xff\xfe'):
            # UTF-16 LE
            with open(file_path, 'r', encoding='utf-16-le') as f:
                return f.read()
        elif first_bytes.startswith(b'\xfe\xff'):
            # UTF-16 BE
            with open(file_path, 'r', encoding='utf-16-be') as f:
                return f.read()
        elif first_bytes.startswith(b'\xef\xbb\xbf'):
            # UTF-8 with BOM
            with open(file_path, 'r', encoding='utf-8-sig') as f:
                return f.read()
    
    # Try common encodings in order
    encodings_to_try = ['utf-8', 'utf-8-sig', 'latin-1', 'cp1252', 'iso-8859-1']
    
    for encoding in encodings_to_try:
        try:
            with open(file_path, 'r', encoding=encoding) as f:
                return f.read()
        except (UnicodeDecodeError, LookupError):
            continue
    
    # Last resort: read with error handling (replace invalid chars)
    with open(file_path, 'r', encoding='utf-8', errors='replace') as f:
        return f.read()


# ============================================================================
# Settings Management
# ============================================================================

class SettingsManager:
    """Manages application settings with save/load functionality"""
    
    SETTINGS_FILE = "config/md_converter_settings.json"
    
    DEFAULT_SETTINGS = {
        # Optional features (Features 2-5)
        "page_breaks_enabled": True,
        "preprocessing_enabled": True,
        "advanced_docx_enabled": True,
        "font_size_options_enabled": True,
        
        # Font size (Feature 5)
        "default_font_size": 10,  # 9, 10, 11, 12, 14
        
        # DOCX settings (Feature 4)
        "docx_narrow_margins": False,
        "docx_language": "de-CH",
        
        # Page break settings (Feature 2)
        "auto_page_break_h2": True,
        "auto_page_break_h3": True,
        "skip_page_break_for": ["einleitung", "zusammenfassung"],
        
        # Auto-open settings
        "autoopen_enabled": True,  # Automatically open PDF/DOCX after generation
    }
    
    def __init__(self):
        self.settings = self.DEFAULT_SETTINGS.copy()
        self.load_settings()
    
    def load_settings(self):
        """Load settings from file"""
        settings_path = Path(self.SETTINGS_FILE)
        if settings_path.exists():
            try:
                with open(settings_path, 'r', encoding='utf-8') as f:
                    loaded = json.load(f)
                    # Merge with defaults to ensure all keys exist
                    self.settings.update(loaded)
            except Exception as e:
                print(f"[WARNING] Failed to load settings: {e}")
                self.settings = self.DEFAULT_SETTINGS.copy()
    
    def save_settings(self):
        """Save settings to file"""
        settings_path = Path(self.SETTINGS_FILE)
        settings_path.parent.mkdir(parents=True, exist_ok=True)
        try:
            with open(settings_path, 'w', encoding='utf-8') as f:
                json.dump(self.settings, f, indent=2, ensure_ascii=False)
        except Exception as e:
            print(f"[ERROR] Failed to save settings: {e}")
    
    def get(self, key: str, default=None):
        """Get a setting value"""
        return self.settings.get(key, default)
    
    def set(self, key: str, value: Any):
        """Set a setting value"""
        self.settings[key] = value
    
    def get_all(self) -> Dict[str, Any]:
        """Get all settings"""
        return self.settings.copy()


# ============================================================================
# Modern UI Styling Constants
# ============================================================================

class UIColors:
    """Modern color palette for consistent UI styling."""
    PRIMARY = "#2563eb"
    PRIMARY_HOVER = "#1d4ed8"
    PRIMARY_LIGHT = "#dbeafe"
    
    SECONDARY = "#64748b"
    SECONDARY_HOVER = "#475569"
    
    SUCCESS = "#16a34a"
    SUCCESS_LIGHT = "#dcfce7"
    SUCCESS_HOVER = "#15803d"
    ERROR = "#dc2626"
    ERROR_LIGHT = "#fee2e2"
    ERROR_HOVER = "#b91c1c"
    WARNING = "#f59e0b"
    WARNING_LIGHT = "#fef3c7"
    
    BG_PRIMARY = "#ffffff"
    BG_SECONDARY = "#f8fafc"
    BG_TERTIARY = "#f1f5f9"
    BORDER = "#e2e8f0"
    BORDER_DARK = "#cbd5e1"
    TEXT_PRIMARY = "#1e293b"
    TEXT_SECONDARY = "#64748b"
    TEXT_MUTED = "#94a3b8"
    
    DROP_ZONE_BG = "#f8fafc"
    DROP_ZONE_BORDER = "#94a3b8"
    DROP_ZONE_ACTIVE = "#dbeafe"
    DROP_ZONE_BORDER_ACTIVE = "#2563eb"


class UIFonts:
    """Font configurations for consistent typography."""
    TITLE = ("Segoe UI", 18, "bold")
    SUBTITLE = ("Segoe UI", 14, "bold")
    HEADING = ("Segoe UI", 12, "bold")
    BODY = ("Segoe UI", 10)
    BODY_BOLD = ("Segoe UI", 10, "bold")
    SMALL = ("Segoe UI", 9)
    SMALL_BOLD = ("Segoe UI", 9, "bold")
    MONO = ("Consolas", 9)
    BUTTON = ("Segoe UI", 10, "bold")
    BUTTON_SMALL = ("Segoe UI", 9)


class UISpacing:
    """Consistent spacing values."""
    XS = 2
    SM = 5
    MD = 10
    LG = 15
    XL = 20
    XXL = 30


# ============================================================================
# Markdown Converter Core Logic
# ============================================================================

class MarkdownConverter:
    """Core markdown conversion logic"""
    
    def __init__(self, css_preset='default', settings_manager=None):
        if not MARKDOWN_AVAILABLE:
            raise ImportError("markdown library is not installed. Please install with: pip install markdown")
        
        self.css_preset = css_preset
        self.settings = settings_manager or SettingsManager()
        self.md = markdown.Markdown(
            extensions=[
                'codehilite',
                'tables', 
                'toc',
                'fenced_code',
                'nl2br',
                'sane_lists'
            ],
            extension_configs={
                'codehilite': {
                    'css_class': 'highlight',
                    'use_pygments': True
                }
            }
        )
    
    def preprocess_markdown(self, md_content: str) -> str:
        """Preprocess markdown content (Feature 3 - optional)"""
        if not self.settings.get('preprocessing_enabled', True):
            return md_content
        
        # Remove horizontal rules (---) before processing
        md_content = re.sub(r'^---+$', '', md_content, flags=re.MULTILINE)
        md_content = re.sub(r'^\*\*\*+$', '', md_content, flags=re.MULTILINE)
        
        # Normalize excessive line breaks (max 2 consecutive empty lines)
        md_content = re.sub(r'\n{4,}', '\n\n\n', md_content)
        
        return md_content
    
    def process_page_breaks(self, md_content: str) -> Tuple[str, List]:
        """Process page breaks in markdown (Feature 2 - optional)"""
        if not self.settings.get('page_breaks_enabled', True):
            return md_content, []
        
        page_break_markers = []
        page_break_pattern = r'<div[^>]*style="[^"]*page-break[^"]*"[^>]*></div>'
        
        def replace_page_break(match):
            marker = f"<!-- PAGEBREAK_{len(page_break_markers)} -->"
            page_break_markers.append(match.group(0))
            return marker
        
        # Replace page breaks with HTML comment placeholders before markdown conversion
        md_content_processed = re.sub(page_break_pattern, replace_page_break, md_content, flags=re.IGNORECASE)
        
        # Also handle simple <!-- PAGEBREAK --> comments
        simple_pagebreak_pattern = r'<!--\s*PAGEBREAK\s*-->'
        simple_pagebreak_count = len(re.findall(simple_pagebreak_pattern, md_content_processed, re.IGNORECASE))
        for i in range(simple_pagebreak_count):
            marker = f"<!-- PAGEBREAK_{len(page_break_markers)} -->"
            page_break_markers.append('<div style="page-break-before: always;"></div>')
            md_content_processed = re.sub(simple_pagebreak_pattern, marker, md_content_processed, count=1, flags=re.IGNORECASE)
        
        return md_content_processed, page_break_markers
        
    def extract_title_from_markdown(self, md_content: str) -> str:
        """Extract title from markdown content (first H1 header)"""
        lines = md_content.strip().split('\n')
        
        for line in lines:
            line = line.strip()
            if line.startswith('# ') and not line.startswith('## '):
                title = line[2:].strip()
                title = re.sub(r'\*\*(.*?)\*\*', r'\1', title)
                title = re.sub(r'\*(.*?)\*', r'\1', title)
                title = re.sub(r'`(.*?)`', r'\1', title)
                title = re.sub(r'[^\w\s\-\.,:;!?()\[\]]+', '', title).strip()
                if title:
                    return title
        
        first_line = lines[0].strip() if lines else ""
        if first_line and not first_line.startswith(('```', '---', '>')):
            title = re.sub(r'[#*`]+', '', first_line).strip()
            if len(title) > 3 and len(title) < 80:
                return title
        
        return "Konvertiertes Dokument"
    
    def process_gemini_citations(self, md_content: str) -> str:
        """Process Gemini citation format: [cite_start]text[cite: numbers]"""
        citation_pattern = r'\[cite_start\](.*?)\[cite:\s*([^\]]+)\]'
        
        def replace_citation(match):
            text = match.group(1)
            return f'<span class="citation">{text}</span>'
        
        return re.sub(citation_pattern, replace_citation, md_content, flags=re.DOTALL)
    
    def markdown_to_html(self, md_content: str) -> str:
        """Convert markdown to styled HTML"""
        if not MARKDOWN_AVAILABLE:
            raise ImportError("markdown library is not installed")
        
        # Preprocess markdown (Feature 3)
        md_content = self.preprocess_markdown(md_content)
        
        # Process page breaks (Feature 2)
        md_content, page_break_markers = self.process_page_breaks(md_content)
        
        document_title = self.extract_title_from_markdown(md_content)
        md_content = self.process_gemini_citations(md_content)
        
        self.md.reset()
        html_body = self.md.convert(md_content)
        
        # Restore page breaks in HTML
        for i, marker in enumerate([f"<!-- PAGEBREAK_{j} -->" for j in range(len(page_break_markers))]):
            if marker in html_body:
                html_body = html_body.replace(marker, page_break_markers[i])
        
        # Add automatic page breaks before numbered H2 headings (Feature 2)
        if self.settings.get('auto_page_break_h2', True):
            h2_pattern = r'<h2[^>]*>(.*?)</h2>'
            h2_matches = []
            for match in re.finditer(h2_pattern, html_body):
                heading_text = match.group(1).strip()
                # Only add page break for main chapters (starting with number and dot)
                skip_words = self.settings.get('skip_page_break_for', ['einleitung', 'zusammenfassung'])
                if re.match(r'^\d+\.', heading_text) and not any(word in heading_text.lower() for word in skip_words):
                    h2_matches.append((match.start(), heading_text))
            
            # Insert page break divs before main chapter H2s
            page_break_html = '<div style="page-break-before: always; height: 0; margin: 0; padding: 0;"></div>'
            for pos, heading_text in reversed(h2_matches):
                check_start = max(0, pos - 150)
                before_text = html_body[check_start:pos]
                if 'page-break' not in before_text.lower():
                    html_body = html_body[:pos] + page_break_html + '\n' + html_body[pos:]
        
        html_body = re.sub(r'<table>(.*?)</table>', r'<div class="table-container"><table>\1</table></div>', html_body, flags=re.DOTALL)
        html_body = re.sub(r'<ul>(.*?)</ul>', r'<div class="list-container"><ul>\1</ul></div>', html_body, flags=re.DOTALL)
        html_body = re.sub(r'<ol>(.*?)</ol>', r'<div class="list-container"><ol>\1</ol></div>', html_body, flags=re.DOTALL)
        
        # Get CSS based on selected preset (use converter's preset if available, otherwise default)
        css_preset = getattr(self, 'css_preset', 'default')
        css_content = self._get_css_preset(css_preset) + self._get_gemini_css()
        
        full_html = f"""<!DOCTYPE html>
<html lang="de">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{document_title}</title>
    {css_content}
</head>
<body>
    <div class="container">
        {html_body}
    </div>
</body>
</html>"""
        
        return full_html
    
    def _get_css_preset(self, preset_name: str = 'default') -> str:
        """Get CSS styling based on preset name"""
        presets = {
            'default': self._get_default_css_content(),
            'modern': self._get_modern_css_content(),
            'classic': self._get_classic_css_content(),
            'dark': self._get_dark_css_content(),
            'professional': self._get_professional_css_content(),
            'minimal': self._get_minimal_css_content(),
        }
        return presets.get(preset_name, presets['default'])
    
    def _get_default_css(self) -> str:
        """Get default CSS styling (legacy method, uses default preset)"""
        return self._get_css_preset('default')
    
    def _get_default_css_content(self) -> str:
        """Get default CSS styling"""
        return """<style>
:root {
    --primary-color: #2563eb;
    --secondary-color: #16a34a;
    --text-color: #1e293b;
    --bg-color: #f8f9fa;
    --card-bg: #ffffff;
}

* {
    box-sizing: border-box;
}

body {
    font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, sans-serif;
    line-height: 1.7;
    color: var(--text-color);
    max-width: 900px;
    margin: 0 auto;
    padding: 40px 20px;
    background: var(--bg-color);
}

.container {
    background: var(--card-bg);
    border-radius: 12px;
    padding: 40px;
    margin: 20px 0;
    box-shadow: 0 2px 8px rgba(0,0,0,0.1);
}

h1, h2, h3, h4, h5, h6 {
    color: var(--text-color);
    margin-top: 2em;
    margin-bottom: 1em;
    font-weight: 700;
    page-break-after: avoid;
    page-break-inside: avoid;
}

h1 { 
    font-size: 2.5em; 
    margin-top: 0;
    border-bottom: 3px solid var(--primary-color);
    padding-bottom: 10px;
}
h2 { 
    font-size: 2em; 
    color: var(--primary-color); 
    border-bottom: 2px solid #e2e8f0;
    padding-bottom: 8px;
}
h2:not(:first-of-type) {
    page-break-before: always;
}
div[style*="page-break"] + h2 {
    page-break-before: auto;
}
h3 { 
    font-size: 1.5em; 
    color: var(--text-color);
}
h4, h5, h6 {
    page-break-after: avoid;
    page-break-inside: avoid;
}

p { 
    margin-bottom: 1.2em; 
    line-height: 1.8;
    orphans: 3;
    widows: 3;
}

code {
    background: #f1f5f9;
    padding: 2px 6px;
    border-radius: 4px;
    font-family: 'Consolas', 'Monaco', 'Courier New', monospace;
    font-size: 0.9em;
    color: #e74c3c;
}

pre {
    background: #2c3e50;
    color: #ecf0f1;
    padding: 20px;
    border-radius: 8px;
    overflow-x: auto;
    margin: 20px 0;
    border: 1px solid #34495e;
    page-break-inside: avoid;
}

pre code {
    background: transparent;
    padding: 0;
    color: inherit;
    color: #ecf0f1;
}

ul, ol {
    margin: 20px 0;
    padding-left: 30px;
    page-break-inside: avoid;
}

li {
    margin-bottom: 8px;
    line-height: 1.6;
    page-break-inside: avoid;
}

table {
    width: 100%;
    border-collapse: collapse;
    margin: 20px 0;
    background: white;
    box-shadow: 0 1px 3px rgba(0,0,0,0.1);
    page-break-inside: avoid;
}

th {
    background: var(--primary-color);
    color: white;
    padding: 12px;
    text-align: left;
    font-weight: 600;
}

td {
    padding: 12px;
    border-bottom: 1px solid #e2e8f0;
}

tr:hover td {
    background: #f8fafc;
}

blockquote {
    border-left: 4px solid var(--primary-color);
    margin: 20px 0;
    padding: 15px 20px;
    background: #f8fafc;
    font-style: italic;
    page-break-inside: avoid;
}

div[style*="page-break"] {
    page-break-after: always !important;
    height: 0;
    margin: 0;
    padding: 0;
}

div[style*="page-break"] + * {
    margin-top: 0;
}

hr {
    display: none;
    visibility: hidden;
    height: 0;
    margin: 0;
    padding: 0;
}

a {
    color: var(--primary-color);
    text-decoration: none;
}

a:hover {
    text-decoration: underline;
}

.table-container {
    overflow-x: auto;
    margin: 20px 0;
}

.list-container {
    margin: 20px 0;
}
</style>"""
    
    def _get_modern_css_content(self) -> str:
        """Get modern, clean CSS styling"""
        return """<style>
:root {
    --primary-color: #6366f1;
    --secondary-color: #10b981;
    --text-color: #111827;
    --bg-color: #ffffff;
    --card-bg: #ffffff;
    --accent-color: #8b5cf6;
}

* {
    box-sizing: border-box;
}

body {
    font-family: 'Inter', -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, sans-serif;
    line-height: 1.75;
    color: var(--text-color);
    max-width: 920px;
    margin: 0 auto;
    padding: 60px 30px;
    background: linear-gradient(135deg, #f5f7fa 0%, #c3cfe2 100%);
}

.container {
    background: var(--card-bg);
    border-radius: 16px;
    padding: 50px;
    margin: 20px 0;
    box-shadow: 0 20px 60px rgba(0,0,0,0.15);
    border: 1px solid rgba(255,255,255,0.8);
}

h1, h2, h3, h4, h5, h6 {
    color: var(--text-color);
    margin-top: 2.5em;
    margin-bottom: 1.2em;
    font-weight: 800;
    letter-spacing: -0.02em;
}

h1 { 
    font-size: 3em; 
    margin-top: 0;
    background: linear-gradient(135deg, var(--primary-color), var(--accent-color));
    -webkit-background-clip: text;
    -webkit-text-fill-color: transparent;
    background-clip: text;
    padding-bottom: 15px;
    border-bottom: 4px solid transparent;
    border-image: linear-gradient(90deg, var(--primary-color), var(--accent-color)) 1;
}
h2 { 
    font-size: 2.2em; 
    color: var(--primary-color);
    border-bottom: 3px solid #e5e7eb;
    padding-bottom: 10px;
}
h3 { 
    font-size: 1.6em; 
    color: var(--text-color);
}

p { 
    margin-bottom: 1.5em; 
    line-height: 1.9;
    font-size: 1.05em;
}

code {
    background: linear-gradient(135deg, #f3f4f6, #e5e7eb);
    padding: 3px 8px;
    border-radius: 6px;
    font-family: 'Fira Code', 'Consolas', 'Monaco', monospace;
    font-size: 0.9em;
    color: #dc2626;
    font-weight: 500;
}

pre {
    background: linear-gradient(135deg, #1e293b, #0f172a);
    color: #e2e8f0;
    padding: 25px;
    border-radius: 12px;
    overflow-x: auto;
    margin: 25px 0;
    border: 1px solid #334155;
    box-shadow: 0 4px 12px rgba(0,0,0,0.2);
}

pre code {
    background: transparent;
    padding: 0;
    color: #e2e8f0;
}

ul, ol {
    margin: 25px 0;
    padding-left: 35px;
}

li {
    margin-bottom: 10px;
    line-height: 1.8;
}

table {
    width: 100%;
    border-collapse: collapse;
    margin: 25px 0;
    background: white;
    box-shadow: 0 4px 12px rgba(0,0,0,0.08);
    border-radius: 12px;
    overflow: hidden;
}

th {
    background: linear-gradient(135deg, var(--primary-color), var(--accent-color));
    color: white;
    padding: 15px;
    text-align: left;
    font-weight: 700;
    text-transform: uppercase;
    font-size: 0.85em;
    letter-spacing: 0.05em;
}

td {
    padding: 15px;
    border-bottom: 1px solid #f3f4f6;
}

tr:hover td {
    background: #f9fafb;
}

tr:last-child td {
    border-bottom: none;
}

blockquote {
    border-left: 5px solid var(--primary-color);
    margin: 25px 0;
    padding: 20px 25px;
    background: linear-gradient(135deg, #f9fafb, #f3f4f6);
    font-style: italic;
    border-radius: 0 8px 8px 0;
    box-shadow: 0 2px 8px rgba(0,0,0,0.05);
}

a {
    color: var(--primary-color);
    text-decoration: none;
    font-weight: 500;
    border-bottom: 2px solid transparent;
    transition: all 0.2s;
}

a:hover {
    border-bottom-color: var(--primary-color);
}

.table-container {
    overflow-x: auto;
    margin: 25px 0;
}

.list-container {
    margin: 25px 0;
}
</style>"""
    
    def _get_classic_css_content(self) -> str:
        """Get classic, traditional CSS styling with serif fonts"""
        return """<style>
:root {
    --primary-color: #1a472a;
    --secondary-color: #8b4513;
    --text-color: #2c2c2c;
    --bg-color: #faf8f3;
    --card-bg: #ffffff;
}

* {
    box-sizing: border-box;
}

body {
    font-family: 'Georgia', 'Times New Roman', serif;
    line-height: 1.8;
    color: var(--text-color);
    max-width: 800px;
    margin: 0 auto;
    padding: 50px 40px;
    background: var(--bg-color);
}

.container {
    background: var(--card-bg);
    border: 2px solid #d4a574;
    border-radius: 4px;
    padding: 50px;
    margin: 20px 0;
    box-shadow: 0 4px 12px rgba(0,0,0,0.1);
}

h1, h2, h3, h4, h5, h6 {
    color: var(--primary-color);
    margin-top: 2em;
    margin-bottom: 1em;
    font-weight: 700;
    font-family: 'Georgia', serif;
}

h1 { 
    font-size: 2.8em; 
    margin-top: 0;
    text-align: center;
    border-bottom: 4px double var(--primary-color);
    padding-bottom: 15px;
    margin-bottom: 30px;
}
h2 { 
    font-size: 2em; 
    border-bottom: 2px solid var(--secondary-color);
    padding-bottom: 8px;
    margin-top: 2.5em;
}
h3 { 
    font-size: 1.5em; 
    color: var(--text-color);
    font-style: italic;
}

p { 
    margin-bottom: 1.5em; 
    line-height: 1.9;
    text-align: justify;
    font-size: 1.1em;
}

code {
    background: #f5f5f5;
    padding: 2px 6px;
    border-radius: 3px;
    font-family: 'Courier New', monospace;
    font-size: 0.95em;
    color: #8b0000;
    border: 1px solid #ddd;
}

pre {
    background: #2f2f2f;
    color: #f5f5f5;
    padding: 20px;
    border-radius: 4px;
    overflow-x: auto;
    margin: 25px 0;
    border: 2px solid #1a1a1a;
    font-family: 'Courier New', monospace;
}

pre code {
    background: transparent;
    padding: 0;
    color: #f5f5f5;
    border: none;
}

ul, ol {
    margin: 25px 0;
    padding-left: 40px;
}

li {
    margin-bottom: 12px;
    line-height: 1.9;
}

table {
    width: 100%;
    border-collapse: collapse;
    margin: 25px 0;
    background: white;
    border: 2px solid #d4a574;
}

th {
    background: var(--primary-color);
    color: white;
    padding: 12px;
    text-align: left;
    font-weight: 700;
    border: 1px solid #1a472a;
}

td {
    padding: 12px;
    border: 1px solid #d4a574;
}

tr:nth-child(even) {
    background: #faf8f3;
}

blockquote {
    border-left: 5px solid var(--secondary-color);
    margin: 25px 0;
    padding: 20px 25px;
    background: #f9f7f2;
    font-style: italic;
    border-radius: 0 4px 4px 0;
}

a {
    color: var(--primary-color);
    text-decoration: underline;
}

a:hover {
    color: var(--secondary-color);
}

.table-container {
    overflow-x: auto;
    margin: 25px 0;
}

.list-container {
    margin: 25px 0;
}
</style>"""
    
    def _get_dark_css_content(self) -> str:
        """Get dark theme CSS styling"""
        return """<style>
:root {
    --primary-color: #60a5fa;
    --secondary-color: #34d399;
    --text-color: #e5e7eb;
    --bg-color: #0f172a;
    --card-bg: #1e293b;
    --border-color: #334155;
}

* {
    box-sizing: border-box;
}

body {
    font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, sans-serif;
    line-height: 1.7;
    color: var(--text-color);
    max-width: 900px;
    margin: 0 auto;
    padding: 40px 20px;
    background: var(--bg-color);
}

.container {
    background: var(--card-bg);
    border-radius: 12px;
    padding: 40px;
    margin: 20px 0;
    box-shadow: 0 8px 32px rgba(0,0,0,0.4);
    border: 1px solid var(--border-color);
}

h1, h2, h3, h4, h5, h6 {
    color: var(--text-color);
    margin-top: 2em;
    margin-bottom: 1em;
    font-weight: 700;
}

h1 { 
    font-size: 2.5em; 
    margin-top: 0;
    border-bottom: 3px solid var(--primary-color);
    padding-bottom: 10px;
    color: var(--primary-color);
}
h2 { 
    font-size: 2em; 
    color: var(--primary-color); 
    border-bottom: 2px solid var(--border-color);
    padding-bottom: 8px;
}
h3 { 
    font-size: 1.5em; 
    color: var(--secondary-color);
}

p { 
    margin-bottom: 1.2em; 
    line-height: 1.8;
    color: #cbd5e1;
}

code {
    background: #0f172a;
    padding: 2px 6px;
    border-radius: 4px;
    font-family: 'Consolas', 'Monaco', 'Courier New', monospace;
    font-size: 0.9em;
    color: var(--secondary-color);
    border: 1px solid var(--border-color);
}

pre {
    background: #0a0e1a;
    color: #e2e8f0;
    padding: 20px;
    border-radius: 8px;
    overflow-x: auto;
    margin: 20px 0;
    border: 1px solid var(--border-color);
    box-shadow: inset 0 2px 8px rgba(0,0,0,0.3);
}

pre code {
    background: transparent;
    padding: 0;
    color: #e2e8f0;
    border: none;
}

ul, ol {
    margin: 20px 0;
    padding-left: 30px;
}

li {
    margin-bottom: 8px;
    line-height: 1.6;
    color: #cbd5e1;
}

table {
    width: 100%;
    border-collapse: collapse;
    margin: 20px 0;
    background: var(--card-bg);
    border: 1px solid var(--border-color);
    border-radius: 8px;
    overflow: hidden;
}

th {
    background: var(--primary-color);
    color: #0f172a;
    padding: 12px;
    text-align: left;
    font-weight: 600;
}

td {
    padding: 12px;
    border-bottom: 1px solid var(--border-color);
    color: #cbd5e1;
}

tr:hover td {
    background: #334155;
}

blockquote {
    border-left: 4px solid var(--primary-color);
    margin: 20px 0;
    padding: 15px 20px;
    background: #1e293b;
    font-style: italic;
    border-radius: 0 8px 8px 0;
}

a {
    color: var(--primary-color);
    text-decoration: none;
}

a:hover {
    color: var(--secondary-color);
    text-decoration: underline;
}

.table-container {
    overflow-x: auto;
    margin: 20px 0;
}

.list-container {
    margin: 20px 0;
}
</style>"""
    
    def _get_professional_css_content(self) -> str:
        """Get professional, business-oriented CSS styling"""
        return """<style>
:root {
    --primary-color: #1e40af;
    --secondary-color: #059669;
    --text-color: #1f2937;
    --bg-color: #f9fafb;
    --card-bg: #ffffff;
    --accent-color: #dc2626;
}

* {
    box-sizing: border-box;
}

body {
    font-family: 'Calibri', 'Arial', 'Helvetica', sans-serif;
    line-height: 1.6;
    color: var(--text-color);
    max-width: 850px;
    margin: 0 auto;
    padding: 40px 30px;
    background: var(--bg-color);
}

.container {
    background: var(--card-bg);
    border-radius: 6px;
    padding: 45px;
    margin: 20px 0;
    box-shadow: 0 2px 8px rgba(0,0,0,0.08);
    border: 1px solid #e5e7eb;
}

h1, h2, h3, h4, h5, h6 {
    color: var(--text-color);
    margin-top: 2em;
    margin-bottom: 1em;
    font-weight: 600;
}

h1 { 
    font-size: 2.2em; 
    margin-top: 0;
    border-bottom: 3px solid var(--primary-color);
    padding-bottom: 12px;
    color: var(--primary-color);
    text-transform: uppercase;
    letter-spacing: 0.05em;
    font-weight: 700;
}
h2 { 
    font-size: 1.8em; 
    color: var(--primary-color); 
    border-bottom: 2px solid #e5e7eb;
    padding-bottom: 8px;
    margin-top: 2em;
}
h3 { 
    font-size: 1.4em; 
    color: var(--text-color);
    font-weight: 600;
}

p { 
    margin-bottom: 1.2em; 
    line-height: 1.7;
    font-size: 1.05em;
}

code {
    background: #f3f4f6;
    padding: 2px 6px;
    border-radius: 3px;
    font-family: 'Courier New', monospace;
    font-size: 0.9em;
    color: var(--accent-color);
    border: 1px solid #e5e7eb;
}

pre {
    background: #1f2937;
    color: #f9fafb;
    padding: 18px;
    border-radius: 4px;
    overflow-x: auto;
    margin: 20px 0;
    border: 1px solid #374151;
    font-family: 'Courier New', monospace;
}

pre code {
    background: transparent;
    padding: 0;
    color: #f9fafb;
    border: none;
}

ul, ol {
    margin: 20px 0;
    padding-left: 30px;
}

li {
    margin-bottom: 8px;
    line-height: 1.7;
}

table {
    width: 100%;
    border-collapse: collapse;
    margin: 20px 0;
    background: white;
    border: 1px solid #e5e7eb;
    box-shadow: 0 1px 3px rgba(0,0,0,0.05);
}

th {
    background: var(--primary-color);
    color: white;
    padding: 12px 15px;
    text-align: left;
    font-weight: 600;
    font-size: 0.95em;
    text-transform: uppercase;
    letter-spacing: 0.03em;
}

td {
    padding: 12px 15px;
    border-bottom: 1px solid #e5e7eb;
}

tr:nth-child(even) {
    background: #f9fafb;
}

tr:hover td {
    background: #f3f4f6;
}

blockquote {
    border-left: 4px solid var(--primary-color);
    margin: 20px 0;
    padding: 15px 20px;
    background: #f9fafb;
    font-style: normal;
    border-radius: 0 4px 4px 0;
}

a {
    color: var(--primary-color);
    text-decoration: none;
    font-weight: 500;
}

a:hover {
    text-decoration: underline;
    color: var(--secondary-color);
}

.table-container {
    overflow-x: auto;
    margin: 20px 0;
}

.list-container {
    margin: 20px 0;
}
</style>"""
    
    def _get_minimal_css_content(self) -> str:
        """Get minimal, clean CSS styling"""
        return """<style>
:root {
    --primary-color: #000000;
    --secondary-color: #666666;
    --text-color: #333333;
    --bg-color: #ffffff;
    --card-bg: #ffffff;
}

* {
    box-sizing: border-box;
}

body {
    font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, sans-serif;
    line-height: 1.6;
    color: var(--text-color);
    max-width: 700px;
    margin: 0 auto;
    padding: 40px 20px;
    background: var(--bg-color);
}

.container {
    background: var(--card-bg);
    padding: 40px;
    margin: 20px 0;
}

h1, h2, h3, h4, h5, h6 {
    color: var(--text-color);
    margin-top: 2em;
    margin-bottom: 1em;
    font-weight: 600;
}

h1 { 
    font-size: 2.2em; 
    margin-top: 0;
    padding-bottom: 10px;
    border-bottom: 1px solid #e5e7eb;
}
h2 { 
    font-size: 1.8em; 
    margin-top: 2em;
    padding-bottom: 8px;
    border-bottom: 1px solid #e5e7eb;
}
h3 { 
    font-size: 1.4em; 
}

p { 
    margin-bottom: 1.2em; 
    line-height: 1.7;
}

code {
    background: #f5f5f5;
    padding: 2px 4px;
    border-radius: 2px;
    font-family: 'Monaco', 'Courier New', monospace;
    font-size: 0.9em;
    color: var(--text-color);
}

pre {
    background: #f5f5f5;
    color: var(--text-color);
    padding: 15px;
    border-radius: 4px;
    overflow-x: auto;
    margin: 20px 0;
    border: 1px solid #e5e7eb;
    font-family: 'Monaco', 'Courier New', monospace;
}

pre code {
    background: transparent;
    padding: 0;
    color: var(--text-color);
}

ul, ol {
    margin: 20px 0;
    padding-left: 25px;
}

li {
    margin-bottom: 6px;
    line-height: 1.6;
}

table {
    width: 100%;
    border-collapse: collapse;
    margin: 20px 0;
    background: white;
}

th {
    background: #f5f5f5;
    color: var(--text-color);
    padding: 10px;
    text-align: left;
    font-weight: 600;
    border-bottom: 2px solid #e5e7eb;
}

td {
    padding: 10px;
    border-bottom: 1px solid #e5e7eb;
}

blockquote {
    border-left: 3px solid #e5e7eb;
    margin: 20px 0;
    padding: 15px 20px;
    background: #f9fafb;
    font-style: italic;
}

a {
    color: var(--text-color);
    text-decoration: underline;
}

a:hover {
    color: var(--secondary-color);
}

.table-container {
    overflow-x: auto;
    margin: 20px 0;
}

.list-container {
    margin: 20px 0;
}
</style>"""
    
    def _get_gemini_css(self) -> str:
        """Get CSS styling for Gemini citations"""
        return """<style>
.citation {
    background: #e3f2fd;
    padding: 8px 12px;
    border-radius: 8px;
    border-left: 4px solid #2196f3;
    margin: 8px 0;
}
</style>"""
    
    def _replace_emojis_for_pdf(self, text: str, keep_emojis: bool = True) -> str:
        """Replace emojis with PDF-safe equivalents or preserve UTF-8 icons"""
        # If keep_emojis is True, preserve all Unicode characters (including UTF-8 icons)
        if keep_emojis:
            return text
        
        # Only replace specific emojis if keep_emojis is False
        emoji_replacements = {
            '✅': '✓', '❌': '✗', '⚠': '!', '⚠️': '!',
            '☐': '[ ]', '□': '[ ]',
            '📝': '', '📄': '', '🚀': '', '💡': '', '🎯': '',
            '📊': '', '📋': '', '🔍': '', '⭐': '', '✨': '',
        }
        
        for emoji, replacement in emoji_replacements.items():
            text = text.replace(emoji, replacement)
        
        # Preserve all Unicode characters (including UTF-8 icons like ✅, ⭐, 🔧, etc.)
        # ReportLab with Unicode-capable fonts (like Segoe UI) can handle these
        return text
    
    def _apply_text_formatting(self, text: str, keep_emojis: bool = True) -> str:
        """Apply text formatting for ReportLab"""
        text = self._replace_emojis_for_pdf(text, keep_emojis=keep_emojis)
        
        citation_pattern = r'\[cite_start\](.*?)\[cite:\s*([^\]]+)\]'
        def replace_citation(match):
            text = match.group(1)
            text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
            text = re.sub(r'\*(.*?)\*', r'\1', text)
            return f'<b>{text}</b>'
        
        text = re.sub(citation_pattern, replace_citation, text, flags=re.DOTALL)
        text = re.sub(r'\*\*(.*?)\*\*', r'<b>\1</b>', text)
        text = re.sub(r'(?<!\*)\*([^*]+)\*(?!\*)', r'<i>\1</i>', text)
        text = re.sub(r'`([^`]+)`', r'<font name="Courier">\1</font>', text)
        text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'<u>\1</u>', text)
        
        text = text.replace('&', '&amp;')
        text = text.replace('<', '&lt;')
        text = text.replace('>', '&gt;')
        
        text = text.replace('&lt;b&gt;', '<b>')
        text = text.replace('&lt;/b&gt;', '</b>')
        text = text.replace('&lt;i&gt;', '<i>')
        text = text.replace('&lt;/i&gt;', '</i>')
        text = text.replace('&lt;u&gt;', '<u>')
        text = text.replace('&lt;/u&gt;', '</u>')
        text = text.replace('&lt;font name="Courier"&gt;', '<font name="Courier">')
        text = text.replace('&lt;/font&gt;', '</font>')
        
        return text
    
    def markdown_to_pdf_reportlab(
        self,
        md_content: str,
        output_path: str,
        orientation: str = 'portrait',
        keep_icons: bool = True,  # Default to True to preserve UTF-8 icons
    ) -> bool:
        """Convert markdown to PDF using ReportLab"""
        if not REPORTLAB_AVAILABLE:
            return False
            
        try:
            base_font = "Helvetica"
            bold_font = "Helvetica-Bold"
            mono_font = "Courier"
            
            try:
                segoe = r"c:\windows\fonts\segoeui.ttf"
                segoe_bold = r"c:\windows\fonts\segoeuib.ttf"
                consola = r"c:\windows\fonts\consola.ttf"
                # Try to register Segoe UI with Unicode support
                if os.path.exists(segoe):
                    try:
                        # Register with Unicode support (subfontIndex=0 enables full Unicode)
                        pdfmetrics.registerFont(TTFont("SegoeUI", segoe, subfontIndex=0))
                        base_font = "SegoeUI"
                    except Exception:
                        # Fallback to standard registration
                        pdfmetrics.registerFont(TTFont("SegoeUI", segoe))
                        base_font = "SegoeUI"
                if os.path.exists(segoe_bold):
                    try:
                        pdfmetrics.registerFont(TTFont("SegoeUI-Bold", segoe_bold, subfontIndex=0))
                        bold_font = "SegoeUI-Bold"
                    except Exception:
                        pdfmetrics.registerFont(TTFont("SegoeUI-Bold", segoe_bold))
                        bold_font = "SegoeUI-Bold"
                if os.path.exists(consola):
                    try:
                        pdfmetrics.registerFont(TTFont("Consolas", consola, subfontIndex=0))
                        mono_font = "Consolas"
                    except Exception:
                        pdfmetrics.registerFont(TTFont("Consolas", consola))
                        mono_font = "Consolas"
            except Exception:
                pass
            
            document_title = self.extract_title_from_markdown(md_content)
            
            if orientation.lower() == 'landscape':
                page_size = landscape(A4)
            else:
                page_size = portrait(A4)
            
            doc = SimpleDocTemplate(
                output_path,
                pagesize=page_size,
                rightMargin=72,
                leftMargin=72,
                topMargin=72,
                bottomMargin=72,
                title=document_title
            )
            
            styles = getSampleStyleSheet()
            
            body_style = ParagraphStyle(
                'Body',
                parent=styles['Normal'],
                fontName=base_font,
                fontSize=11,
                leading=14,
                textColor=colors.HexColor('#2c3e50'),
                alignment=TA_JUSTIFY,  # Justified text for better readability
            )
            
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=24,
                spaceAfter=30,
                textColor=colors.HexColor('#2c3e50'),
                fontName=bold_font
            )
            
            heading_style = ParagraphStyle(
                'CustomHeading',
                parent=styles['Heading2'],
                fontSize=16,
                spaceAfter=12,
                textColor=colors.HexColor('#2c3e50'),
                fontName=bold_font
            )
            
            code_style = ParagraphStyle(
                'CodeStyle',
                parent=styles['Code'],
                fontSize=9,
                fontName=mono_font,
                backgroundColor=colors.HexColor('#f4f4f4'),
                borderColor=colors.HexColor('#e9ecef'),
                borderWidth=1,
                borderPadding=10,
                leftIndent=20,
                rightIndent=20,
                leading=12,
                spaceAfter=12
            )
            
            story = []
            lines = md_content.split('\n')
            in_code_block = False
            code_lines = []
            i = 0
            
            while i < len(lines):
                line = lines[i]
                line_stripped = line.strip()
                
                if line_stripped.startswith('```'):
                    if in_code_block:
                        if code_lines:
                            code_text = '\n'.join(code_lines)
                            safe_code = (
                                code_text
                                .replace('&', '&amp;')
                                .replace('<', '&lt;')
                                .replace('>', '&gt;')
                            )
                            story.append(Preformatted(safe_code, code_style))
                            story.append(Spacer(1, 12))
                        code_lines = []
                        in_code_block = False
                    else:
                        in_code_block = True
                        code_lines = []
                    i += 1
                    continue
                
                if in_code_block:
                    code_lines.append(line)
                    i += 1
                    continue
                
                if line_stripped.startswith('#'):
                    level = len(line_stripped) - len(line_stripped.lstrip('#'))
                    text = line_stripped.lstrip('#').strip()
                    text = self._apply_text_formatting(text, keep_emojis=keep_icons)
                    
                    if level == 1:
                        story.append(Paragraph(text, title_style))
                    elif level == 2:
                        story.append(Paragraph(text, heading_style))
                    else:
                        story.append(Paragraph(text, heading_style))
                    story.append(Spacer(1, 6))
                
                elif line_stripped.startswith(('- ', '* ', '+ ')):
                    text = line_stripped[2:].strip()
                    text = self._apply_text_formatting(text, keep_emojis=keep_icons)
                    story.append(Paragraph(f"• {text}", body_style))
                
                elif re.match(r'^\d+\.\s', line_stripped):
                    match = re.match(r'^(\d+)\.\s(.+)', line_stripped)
                    if match:
                        num, text = match.groups()
                        text = self._apply_text_formatting(text, keep_emojis=keep_icons)
                        story.append(Paragraph(f"{num}. {text}", body_style))
                
                # Tables - improved detection (must start with |)
                elif line_stripped.startswith('|') and '|' in line_stripped:
                    table_rows = []
                    j = i
                    
                    # Collect all consecutive table rows
                    while j < len(lines) and lines[j].strip().startswith('|'):
                        row_line = lines[j].strip()
                        # Split and clean cells
                        cells = [cell.strip() for cell in row_line.split('|')[1:-1]]
                        # Skip separator rows (only contains dashes, colons, spaces, and pipes)
                        if cells and not all(re.match(r'^[\s\-:]+$', c) for c in cells):
                            table_rows.append(cells)
                        j += 1
                    
                    if len(table_rows) > 0:
                        table_data = []
                        for row_idx, row in enumerate(table_rows):
                            if row:  # Only process non-empty rows
                                formatted_cells = [
                                    self._apply_text_formatting(cell, keep_emojis=keep_icons)
                                    for cell in row
                                ]
                                # First row is header
                                if row_idx == 0:
                                    table_data.append([Paragraph(c, heading_style) for c in formatted_cells])
                                else:
                                    table_data.append([Paragraph(c, body_style) for c in formatted_cells])
                        
                        if table_data:
                            num_cols = max(len(r) for r in table_data) if table_data else 1
                            col_widths = [doc.width / num_cols] * num_cols
                            t = Table(table_data, colWidths=col_widths, hAlign='LEFT', repeatRows=1)
                            t.setStyle(TableStyle([
                                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#2563eb')),
                                ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
                                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                                ('FONTNAME', (0, 0), (-1, 0), bold_font),
                                ('FONTSIZE', (0, 0), (-1, 0), 10),
                                ('FONTSIZE', (0, 1), (-1, -1), 10),
                                ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
                                ('TOPPADDING', (0, 0), (-1, -1), 8),
                                ('GRID', (0, 0), (-1, -1), 1, colors.HexColor('#d0d7de')),
                                ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#f8f9fa')]),
                                ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                            ]))
                            story.append(t)
                            story.append(Spacer(1, 12))
                    
                    i = j - 1  # -1 because outer loop will increment
                
                elif line_stripped:
                    text = self._apply_text_formatting(line_stripped, keep_emojis=keep_icons)
                    story.append(Paragraph(text, body_style))
                    story.append(Spacer(1, 6))
                else:
                    story.append(Spacer(1, 6))
                
                i += 1
            
            doc.build(story)
            return True
            
        except Exception as e:
            print(f"ReportLab PDF generation failed: {str(e)}")
            return False
    
    def markdown_to_pdf_weasyprint(self, md_content: str, output_path: str, orientation: str = 'portrait') -> bool:
        """Convert Markdown to PDF using WeasyPrint (Feature 6)"""
        if not _try_import_weasyprint():
            return False
        
        try:
            # Preprocess markdown
            md_content = self.preprocess_markdown(md_content)
            
            # Process page breaks
            md_content, page_break_markers = self.process_page_breaks(md_content)
            
            # Convert markdown to HTML
            html_content = self.markdown_to_html(md_content)
            
            # Add @page CSS for orientation
            page_size = "A4 landscape" if orientation.lower() == 'landscape' else "A4"
            page_css = f"""
            @page {{
                size: {page_size};
                margin: 2cm;
            }}
            """
            
            # Inject page CSS into HTML
            html_content = html_content.replace('</style>', f'{page_css}</style>', 1)
            
            # Convert HTML to PDF using WeasyPrint
            HTML = weasyprint['HTML']
            FontConfiguration = weasyprint['FontConfiguration']
            
            font_config = FontConfiguration()
            HTML(string=html_content).write_pdf(output_path, font_config=font_config)
            return True
        except Exception as e:
            print(f"[ERROR] WeasyPrint conversion failed: {e}")
            return False
    
    def _find_chromium_for_pdf(self) -> Optional[str]:
        """Find a Chromium-based browser executable"""
        candidates = [
            r"c:\program files (x86)\microsoft\edge\application\msedge.exe",
            r"c:\program files\microsoft\edge\application\msedge.exe",
            r"c:\program files\google\chrome\application\chrome.exe",
            r"c:\program files (x86)\google\chrome\application\chrome.exe",
        ]
        for p in candidates:
            if os.path.exists(p):
                return p
        return None
    
    def html_to_pdf_browser(
        self,
        html_content: str,
        output_path: str,
        orientation: str = "portrait",
        scale: float = 1.0,
        replace_icons: bool = False,
    ) -> bool:
        """Convert HTML to PDF using browser"""
        exe = self._find_chromium_for_pdf()
        if not exe:
            return False
        
        try:
            if replace_icons:
                html_content = self._replace_icons_for_print(html_content)
            
            temp_dir = tempfile.gettempdir()
            html_path = os.path.join(temp_dir, "md_converter_print.html")
            with open(html_path, "w", encoding="utf-8") as f:
                f.write(html_content)
            
            file_url = Path(html_path).absolute().as_uri()
            
            out_dir = os.path.dirname(os.path.abspath(output_path))
            if out_dir and not os.path.exists(out_dir):
                os.makedirs(out_dir, exist_ok=True)
            
            pdf_out = os.path.abspath(output_path).replace("\\", "/")
            user_data_dir = tempfile.mkdtemp(prefix="md_converter_chromium_profile_")
            
            base_args = [
                exe,
                "--headless=new",
                "--disable-gpu",
                "--no-first-run",
                "--no-default-browser-check",
                "--disable-extensions",
                "--disable-dev-shm-usage",
                "--allow-file-access-from-files",
                "--disable-web-security",
                "--run-all-compositor-stages-before-draw",
                "--virtual-time-budget=8000",
                "--no-pdf-header-footer",
                f"--user-data-dir={user_data_dir}",
                f"--print-to-pdf={pdf_out}",
                file_url,
            ]
            
            try:
                r = subprocess.run(base_args, capture_output=True, text=True, timeout=180)
                if r.returncode == 0 and os.path.exists(output_path):
                    size = os.path.getsize(output_path)
                    if size > 20_000:
                        return True
                return False
            finally:
                try:
                    shutil.rmtree(user_data_dir, ignore_errors=True)
                except Exception:
                    pass
        except Exception:
            return False
    
    def _replace_icons_for_print(self, s: str) -> str:
        """Replace emojis with PDF-safe equivalents"""
        if not s:
            return s
        replacements = {
            "✅": "✓", "❌": "✗", "⚠": "!", "⚠️": "!",
            "☐": "[ ]", "□": "[ ]",
            "✨": "", "📍": "", "📄": "", "📝": "",
        }
        for k, v in replacements.items():
            s = s.replace(k, v)
        return s
    
    def _normalize_anchor_for_docx(self, anchor: str) -> str:
        """Make anchor names safe for Word bookmarks (Feature 4)"""
        safe = re.sub(r'[^A-Za-z0-9_-]+', '-', anchor.strip())
        if not safe or not safe[0].isalpha():
            safe = f"a-{safe}"
        return safe
    
    def _extract_heading_text_and_anchor(self, text: str):
        """Extract visible heading text and optional {#anchor} (Feature 4)"""
        match = re.match(r'^(.*)\s*\{#([A-Za-z0-9_-]+)\}\s*$', text)
        if match:
            return match.group(1).strip(), match.group(2).strip()
        return text.strip(), None
    
    def _add_bookmark(self, paragraph, anchor: str, bookmark_id: int):
        """Add a Word bookmark to the given paragraph (Feature 4)"""
        if not DOCX_AVAILABLE:
            return
        try:
            from docx.oxml import OxmlElement
            from docx.oxml.ns import qn
            
            start = OxmlElement('w:bookmarkStart')
            start.set(qn('w:id'), str(bookmark_id))
            start.set(qn('w:name'), anchor)
            end = OxmlElement('w:bookmarkEnd')
            end.set(qn('w:id'), str(bookmark_id))
            
            paragraph._p.insert(0, start)
            paragraph._p.append(end)
        except Exception:
            pass
    
    def _add_anchor_hyperlink(self, paragraph, text: str, anchor: str):
        """Add internal hyperlink to a bookmark (Feature 4)"""
        if not DOCX_AVAILABLE:
            return
        try:
            from docx.oxml import OxmlElement
            from docx.oxml.ns import qn
            
            hyperlink = OxmlElement('w:hyperlink')
            hyperlink.set(qn('w:anchor'), anchor)
            
            run = OxmlElement('w:r')
            r_pr = OxmlElement('w:rPr')
            r_style = OxmlElement('w:rStyle')
            r_style.set(qn('w:val'), 'Hyperlink')
            r_pr.append(r_style)
            run.append(r_pr)
            
            text_elem = OxmlElement('w:t')
            text_elem.text = text
            run.append(text_elem)
            
            hyperlink.append(run)
            paragraph._p.append(hyperlink)
        except Exception:
            pass
    
    def _add_external_hyperlink(self, paragraph, text: str, url: str):
        """Add external hyperlink to a URL (Feature 4)"""
        if not DOCX_AVAILABLE:
            return
        try:
            from docx.oxml import OxmlElement
            from docx.oxml.ns import qn
            import docx.opc.constants as constants
            
            part = paragraph.part
            r_id = part.relate_to(url, constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
            
            hyperlink = OxmlElement('w:hyperlink')
            hyperlink.set(qn('r:id'), r_id)
            
            run = OxmlElement('w:r')
            r_pr = OxmlElement('w:rPr')
            r_style = OxmlElement('w:rStyle')
            r_style.set(qn('w:val'), 'Hyperlink')
            r_pr.append(r_style)
            run.append(r_pr)
            
            text_elem = OxmlElement('w:t')
            text_elem.text = text
            run.append(text_elem)
            
            hyperlink.append(run)
            paragraph._p.append(hyperlink)
        except Exception:
            pass
    
    def _set_default_language(self, doc, lang_code="de-CH"):
        """Set default proofing language for the document (Feature 4)"""
        if not DOCX_AVAILABLE:
            return
        try:
            from docx.oxml.ns import qn
            style = doc.styles['Normal']
            rpr = style._element.get_or_add_rPr()
            lang = rpr.get_or_add_lang()
            lang.set(qn('w:val'), lang_code)
        except Exception:
            pass
    
    def _add_formatted_text(self, paragraph, text, anchor_map=None):
        """Add inline markdown formatting and links (Feature 4)"""
        if not DOCX_AVAILABLE:
            paragraph.add_run(text)
            return
        
        anchor_map = anchor_map or {}
        
        def add_run(content, style=None):
            run = paragraph.add_run(content)
            if style == 'bold':
                run.bold = True
            elif style == 'italic':
                run.italic = True
            elif style == 'code':
                run.font.name = 'Courier New'
        
        # Match links: [[text]](url) and [text](url)
        double_bracket_pattern = re.compile(r'\[\[([^\]]+)\]\]\(([^)]+)\)')
        single_bracket_pattern = re.compile(r'(?<!\[)\[([^\]]+)\]\(([^)]+)\)(?!\])')
        
        matches = []
        for match in double_bracket_pattern.finditer(text):
            matches.append((match.start(), match.end(), match.group(1), match.group(2), True))
        for match in single_bracket_pattern.finditer(text):
            overlap = False
            for d_start, d_end, _, _, _ in matches:
                if not (match.end() <= d_start or match.start() >= d_end):
                    overlap = True
                    break
            if not overlap:
                matches.append((match.start(), match.end(), match.group(1), match.group(2), False))
        
        matches.sort(key=lambda x: x[0])
        
        idx = 0
        for match_start, match_end, link_text, link_target, is_double_bracket in matches:
            if match_start > idx:
                self._add_formatted_text(paragraph, text[idx:match_start], anchor_map=anchor_map)
            
            if is_double_bracket:
                link_text = f"[{link_text}]"
            
            if link_target.startswith('#'):
                link_anchor = link_target[1:]
                target = anchor_map.get(link_anchor)
                if target:
                    self._add_anchor_hyperlink(paragraph, link_text, target)
                else:
                    add_run(link_text)
            elif link_target.startswith('http://') or link_target.startswith('https://'):
                self._add_external_hyperlink(paragraph, link_text, link_target)
            else:
                add_run(link_text)
            idx = match_end
        
        remaining = text[idx:]
        if not remaining:
            return
        
        parts = re.split(r'(\*\*.*?\*\*|_.*?_|`.*?`)', remaining)
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                add_run(part[2:-2], 'bold')
            elif part.startswith('_') and part.endswith('_') and len(part) > 2:
                add_run(part[1:-1], 'italic')
            elif part.startswith('`') and part.endswith('`'):
                add_run(part[1:-1], 'code')
            else:
                add_run(part)
    
    def markdown_to_docx(self, md_content: str, output_path: str) -> bool:
        """Convert markdown to DOCX document (Feature 4: Advanced DOCX)"""
        if not DOCX_AVAILABLE:
            messagebox.showerror("Error", "python-docx not available. Cannot export to DOCX.")
            return False
        
        use_advanced = self.settings.get('advanced_docx_enabled', True)
        font_size = self.settings.get('default_font_size', 10) if self.settings.get('font_size_options_enabled', True) else None
        narrow_margins = self.settings.get('docx_narrow_margins', False)
        language = self.settings.get('docx_language', 'de-CH')
        
        try:
            # Preprocess markdown
            md_content = self.preprocess_markdown(md_content)
            
            doc = Document()
            document_title = self.extract_title_from_markdown(md_content)
            doc.core_properties.title = document_title
            doc.core_properties.author = "Markdown Converter"
            
            # Set default language
            if use_advanced:
                self._set_default_language(doc, language)
            
            # Set margins
            if narrow_margins:
                sections = doc.sections
                for section in sections:
                    section.top_margin = Inches(0.5)
                    section.bottom_margin = Inches(0.5)
                    section.left_margin = Inches(0.5)
                    section.right_margin = Inches(0.5)
            
            # Set font size
            if font_size is not None:
                style = doc.styles['Normal']
                style.font.size = Pt(font_size)
                # Adjust heading sizes proportionally
                heading_multipliers = {
                    'Heading 1': 1.5,
                    'Heading 2': 1.3,
                    'Heading 3': 1.15,
                    'Heading 4': 1.1,
                    'Heading 5': 1.1,
                    'Heading 6': 1.1,
                }
                for heading_name, multiplier in heading_multipliers.items():
                    heading_style = doc.styles[heading_name]
                    heading_style.font.size = Pt(int(font_size * multiplier))
            
            md_content = self.process_gemini_citations(md_content)
            
            lines = md_content.split('\n')
            in_code_block = False
            code_lines = []
            bookmark_id = 0
            anchor_map = {}
            previous_was_content = False
            needs_page_break = False
            
            i = 0
            while i < len(lines):
                line = lines[i]
                line_stripped = line.strip()
                
                # Handle page breaks
                if 'page-break' in line_stripped.lower() and 'div' in line_stripped.lower():
                    needs_page_break = True
                    i += 1
                    continue
                if re.match(r'<!--\s*PAGEBREAK\s*-->', line_stripped, re.IGNORECASE):
                    needs_page_break = True
                    i += 1
                    continue
                
                if line_stripped.startswith('```'):
                    if in_code_block:
                        if code_lines:
                            code_text = '\n'.join(code_lines)
                            if use_advanced:
                                p = doc.add_paragraph()
                                run = p.add_run(code_text)
                                run.font.name = 'Courier New'
                            else:
                                p = doc.add_paragraph(code_text, style='Code')
                        code_lines = []
                        in_code_block = False
                    else:
                        in_code_block = True
                        code_lines = []
                    i += 1
                    continue
                
                if in_code_block:
                    code_lines.append(line)
                    i += 1
                    continue
                
                if line_stripped.startswith('#'):
                    level = len(line_stripped) - len(line_stripped.lstrip('#'))
                    heading_text_raw = line_stripped.lstrip('#').strip()
                    
                    if use_advanced:
                        heading_text, heading_anchor = self._extract_heading_text_and_anchor(heading_text_raw)
                    else:
                        heading_text = heading_text_raw
                        heading_anchor = None
                    
                    lvl = min(level, 6)
                    p = doc.add_heading(heading_text, level=lvl)
                    
                    # Add bookmark if anchor exists
                    if use_advanced and heading_anchor:
                        normalized = self._normalize_anchor_for_docx(heading_anchor)
                        anchor_map[heading_anchor] = normalized
                        self._add_bookmark(p, normalized, bookmark_id)
                        bookmark_id += 1
                    
                    # Add page break for numbered headings if enabled
                    if use_advanced and needs_page_break:
                        try:
                            from docx.enum.text import WD_BREAK
                            para = doc.add_paragraph()
                            run = para.add_run()
                            run.add_break(WD_BREAK.PAGE)
                        except Exception:
                            pass
                        needs_page_break = False
                    
                    previous_was_content = False
                
                elif line_stripped.startswith(('- ', '* ', '+ ')):
                    text = line_stripped[2:].strip()
                    p = doc.add_paragraph(style='List Bullet')
                    if use_advanced:
                        self._add_formatted_text(p, text, anchor_map=anchor_map)
                    else:
                        p.add_run(text)
                    previous_was_content = True
                
                elif re.match(r'^\d+\.\s', line_stripped):
                    text = re.sub(r'^\d+\.\s', '', line_stripped)
                    p = doc.add_paragraph(style='List Number')
                    if use_advanced:
                        self._add_formatted_text(p, text, anchor_map=anchor_map)
                    else:
                        p.add_run(text)
                    previous_was_content = True
                
                elif line_stripped:
                    p = doc.add_paragraph()
                    if use_advanced:
                        self._add_formatted_text(p, line_stripped, anchor_map=anchor_map)
                    else:
                        p.add_run(line_stripped)
                    previous_was_content = True
                
                i += 1
            
            doc.save(output_path)
            return True
            
        except Exception as e:
            messagebox.showerror("DOCX Export Error", f"Failed to create DOCX: {str(e)}")
            return False


# ============================================================================
# GUI Application
# ============================================================================

class MarkdownConverterGUI:
    """GUI application for markdown conversion"""
    
    def __init__(self):
        if HAS_DND:
            self.root = TkinterDnD.Tk()
        else:
            self.root = tk.Tk()
        
        self.root.title("Markdown to PDF/DOCX Converter")
        self.position_window()
        
        # Initialize CSS preset before creating converter
        self.css_preset = "default"  # Current CSS preset
        
        # Initialize settings manager first
        self.settings_manager = SettingsManager()
        
        # Check if markdown is available before creating converter
        if not MARKDOWN_AVAILABLE:
            self.show_missing_dependency_error()
            self.converter = None
        else:
            try:
                # Create converter with settings
                self.converter = MarkdownConverter(css_preset=self.css_preset, settings_manager=self.settings_manager)
            except ImportError as e:
                self.show_missing_dependency_error(str(e))
                self.converter = None
        
        self.current_html = ""
        self.browser_preview_path: Optional[str] = None
        self.current_file_path: Optional[str] = None  # Track current file for save dialog
        self.preview_html = None  # HTML browser widget
        self.preview_text = None  # Text preview widget (fallback)
        self.use_html_preview = False  # Whether to use HTML browser or text preview
        
        self.setup_ui()
        if HAS_DND:
            self.setup_drag_drop()
        
        # Save settings on exit
        self.root.protocol("WM_DELETE_WINDOW", self.on_closing)
    
    def show_missing_dependency_error(self, error_msg=None):
        """Show error dialog for missing dependencies"""
        msg = "The 'markdown' library is required but not installed.\n\n"
        msg += "Please install it with:\n"
        msg += "  pip install markdown\n\n"
        if error_msg:
            msg += f"Error: {error_msg}\n\n"
        msg += "The tool will start but conversion features will be disabled."
        
        messagebox.showerror("Missing Dependency", msg)
    
    def position_window(self):
        """Position window in the area below the launcher"""
        x = int(os.environ.get('TOOL_WINDOW_X', 100))
        y = int(os.environ.get('TOOL_WINDOW_Y', 100))
        width = int(os.environ.get('TOOL_WINDOW_WIDTH', 1200))
        height = int(os.environ.get('TOOL_WINDOW_HEIGHT', 800))
        
        # Ensure minimum size (match launcher minimum width)
        width = max(width, 1280)  # Match launcher minimum width (1280px)
        height = max(height, 600)
        
        self.root.geometry(f"{width}x{height}+{x}+{y}")
        
        # Set minimum size to match launcher
        self.root.minsize(1280, 600)
    
    def setup_ui(self):
        """Set up the user interface"""
        self.root.configure(bg=UIColors.BG_SECONDARY)
        
        # Main container
        main_frame = tk.Frame(self.root, bg=UIColors.BG_SECONDARY)
        main_frame.pack(fill='both', expand=True, padx=UISpacing.MD, pady=UISpacing.MD)
        
        # Header
        header_frame = tk.Frame(main_frame, bg=UIColors.BG_PRIMARY, pady=UISpacing.MD)
        header_frame.pack(fill='x', pady=(0, UISpacing.MD))
        
        title_label = tk.Label(
            header_frame,
            text="📄 Markdown to PDF/DOCX Converter",
            font=UIFonts.TITLE,
            bg=UIColors.BG_PRIMARY,
            fg=UIColors.PRIMARY
        )
        title_label.pack()
        
        # Toolbar
        toolbar = tk.Frame(main_frame, bg=UIColors.BG_PRIMARY)
        toolbar.pack(fill='x', pady=(0, UISpacing.MD))
        
        self.create_rounded_button(toolbar, "📁 Open File", self.open_file, "primary").pack(side='left', padx=UISpacing.XS)
        self.create_rounded_button(toolbar, "💾 Save as PDF", self.save_pdf, "success").pack(side='left', padx=UISpacing.XS)
        self.create_rounded_button(toolbar, "📄 Save as DOCX", self.save_docx, "primary").pack(side='left', padx=UISpacing.XS)
        self.create_rounded_button(toolbar, "🌐 Browser Preview", self.open_in_browser, "secondary").pack(side='left', padx=UISpacing.XS)
        self.create_rounded_button(toolbar, "⚙️ Settings", self.show_settings_dialog, "secondary").pack(side='left', padx=UISpacing.XS)
        
        # PDF settings
        pdf_settings = tk.Frame(main_frame, bg=UIColors.BG_PRIMARY)
        pdf_settings.pack(fill='x', pady=(0, UISpacing.MD))
        
        tk.Label(pdf_settings, text="PDF Engine:", bg=UIColors.BG_PRIMARY, font=UIFonts.SMALL).pack(side='left', padx=UISpacing.SM)
        self.pdf_engine = tk.StringVar(value="browser")  # Browser engine as default
        tk.Radiobutton(pdf_settings, text="Browser", variable=self.pdf_engine, value="browser", bg=UIColors.BG_PRIMARY).pack(side='left', padx=UISpacing.XS)
        if _try_import_weasyprint():
            tk.Radiobutton(pdf_settings, text="WeasyPrint", variable=self.pdf_engine, value="weasyprint", bg=UIColors.BG_PRIMARY).pack(side='left', padx=UISpacing.XS)
        tk.Radiobutton(pdf_settings, text="ReportLab", variable=self.pdf_engine, value="reportlab", bg=UIColors.BG_PRIMARY).pack(side='left', padx=UISpacing.XS)
        
        tk.Label(pdf_settings, text="Layout:", bg=UIColors.BG_PRIMARY, font=UIFonts.SMALL).pack(side='left', padx=(UISpacing.LG, UISpacing.SM))
        self.pdf_orientation = tk.StringVar(value="portrait")
        tk.Radiobutton(pdf_settings, text="Portrait", variable=self.pdf_orientation, value="portrait", bg=UIColors.BG_PRIMARY).pack(side='left', padx=UISpacing.XS)
        tk.Radiobutton(pdf_settings, text="Landscape", variable=self.pdf_orientation, value="landscape", bg=UIColors.BG_PRIMARY).pack(side='left', padx=UISpacing.XS)
        
        # HTML Style Preset selector
        style_settings = tk.Frame(main_frame, bg=UIColors.BG_PRIMARY)
        style_settings.pack(fill='x', pady=(0, UISpacing.MD))
        
        tk.Label(style_settings, text="HTML Style:", bg=UIColors.BG_PRIMARY, font=UIFonts.SMALL).pack(side='left', padx=UISpacing.SM)
        self.css_preset_var = tk.StringVar(value="default")
        self.css_preset_var.trace('w', lambda *args: self.on_css_preset_change())
        
        presets = [
            ("Default", "default"),
            ("Modern", "modern"),
            ("Classic", "classic"),
            ("Dark", "dark"),
            ("Professional", "professional"),
            ("Minimal", "minimal")
        ]
        
        for label, value in presets:
            tk.Radiobutton(
                style_settings, 
                text=label, 
                variable=self.css_preset_var, 
                value=value, 
                bg=UIColors.BG_PRIMARY,
                font=UIFonts.SMALL
            ).pack(side='left', padx=UISpacing.XS)
        
        # Content area
        content_frame = tk.Frame(main_frame, bg=UIColors.BG_SECONDARY)
        content_frame.pack(fill='both', expand=True)
        
        # Left: Input
        left_frame = tk.LabelFrame(content_frame, text=" Markdown Input ", font=UIFonts.HEADING, bg=UIColors.BG_PRIMARY, fg=UIColors.TEXT_PRIMARY)
        left_frame.pack(side='left', fill='both', expand=True, padx=(0, UISpacing.SM))
        
        self.text_input = scrolledtext.ScrolledText(
            left_frame,
            wrap=tk.WORD,
            font=UIFonts.MONO,
            bg=UIColors.BG_PRIMARY,
            fg=UIColors.TEXT_PRIMARY,
            insertbackground=UIColors.PRIMARY
        )
        self.text_input.pack(fill='both', expand=True, padx=UISpacing.SM, pady=UISpacing.SM)
        self.text_input.bind('<KeyRelease>', self.on_text_change)
        
        # Right: Preview
        right_frame = tk.LabelFrame(content_frame, text=" HTML Preview ", font=UIFonts.HEADING, bg=UIColors.BG_PRIMARY, fg=UIColors.TEXT_PRIMARY)
        right_frame.pack(side='right', fill='both', expand=True, padx=(UISpacing.SM, 0))
        
        # Use embedded HTML browser if available, otherwise fall back to text preview
        if TKINTERWEB_AVAILABLE:
            try:
                # Create embedded HTML browser widget
                self.preview_html = HtmlFrame(right_frame, messages_enabled=False)
                self.preview_html.pack(fill='both', expand=True, padx=UISpacing.SM, pady=UISpacing.SM)
                self.preview_text = None  # Not used when HTML browser is available
                self.use_html_preview = True
            except Exception as e:
                print(f"[WARNING] Failed to create HTML browser widget: {e}")
                # Fall back to text preview
                self.preview_text = scrolledtext.ScrolledText(
                    right_frame,
                    wrap=tk.WORD,
                    font=UIFonts.BODY,
                    bg=UIColors.BG_PRIMARY,
                    fg=UIColors.TEXT_PRIMARY,
                    state='disabled'
                )
                self.preview_text.pack(fill='both', expand=True, padx=UISpacing.SM, pady=UISpacing.SM)
                self.preview_html = None
                self.use_html_preview = False
        else:
            # Fall back to text preview
            self.preview_text = scrolledtext.ScrolledText(
                right_frame,
                wrap=tk.WORD,
                font=UIFonts.BODY,
                bg=UIColors.BG_PRIMARY,
                fg=UIColors.TEXT_PRIMARY,
                state='disabled'
            )
            self.preview_text.pack(fill='both', expand=True, padx=UISpacing.SM, pady=UISpacing.SM)
            self.preview_html = None
            self.use_html_preview = False
        
        # Status bar
        self.status_var = tk.StringVar(value="Ready - Drag .md files here or paste markdown content")
        status_bar = tk.Label(
            main_frame,
            textvariable=self.status_var,
            bg=UIColors.BG_TERTIARY,
            fg=UIColors.TEXT_SECONDARY,
            font=UIFonts.SMALL,
            anchor='w',
            padx=UISpacing.SM,
            pady=UISpacing.XS
        )
        status_bar.pack(fill='x', pady=(UISpacing.MD, 0))
        
        self.load_sample_content()
    
    def create_rounded_button(self, parent, text, command, style="primary"):
        """Create a styled button"""
        colors_map = {
            "primary": (UIColors.PRIMARY, UIColors.PRIMARY_HOVER, "#ffffff"),
            "secondary": (UIColors.BG_TERTIARY, UIColors.BORDER, UIColors.TEXT_PRIMARY),
            "success": (UIColors.SUCCESS, UIColors.SUCCESS_HOVER, "#ffffff"),
        }
        
        bg, hover_bg, fg = colors_map.get(style, colors_map["primary"])
        
        btn = tk.Button(
            parent,
            text=text,
            command=command,
            font=UIFonts.BUTTON,
            bg=bg,
            fg=fg,
            activebackground=hover_bg,
            activeforeground=fg,
            relief="flat",
            cursor="hand2",
            padx=UISpacing.MD,
            pady=UISpacing.SM,
            bd=0,
            highlightthickness=0,
        )
        
        def on_enter(e):
            btn.config(bg=hover_bg)
        def on_leave(e):
            btn.config(bg=bg)
        
        btn.bind("<Enter>", on_enter)
        btn.bind("<Leave>", on_leave)
        
        return btn
    
    def setup_drag_drop(self):
        """Set up drag and drop functionality"""
        if HAS_DND:
            self.root.drop_target_register(DND_FILES)
            self.root.dnd_bind('<<Drop>>', self.on_drop)
            self.text_input.drop_target_register(DND_FILES)
            self.text_input.dnd_bind('<<Drop>>', self.on_drop)
    
    def on_drop(self, event):
        """Handle file drop events"""
        files = self.parse_dropped_files(event.data)
        for file_path in files:
            if file_path.lower().endswith('.md'):
                self.load_file(file_path)
                break
        return "break"
    
    def parse_dropped_files(self, data):
        """Parse dropped file paths"""
        files = []
        if '{' in data:
            import re
            files = re.findall(r'\{([^}]+)\}', data)
            remaining = re.sub(r'\{[^}]+\}', '', data).strip()
            if remaining:
                files.extend(remaining.split())
        else:
            files = data.split()
        return files
    
    def load_sample_content(self):
        """Load sample markdown content"""
        sample = """# Markdown to PDF/DOCX Converter

Willkommen beim **Markdown-Converter** mit moderner Formatierung!

## Features

- **PDF Export** mit ReportLab oder Browser-Engine
- **DOCX Export** für Word-Kompatibilität
- **Live Preview** im Browser
- **Drag & Drop** Unterstützung

## Beispiel-Text

Dies ist ein Beispiel für **fettgedruckten** und *kursiven* Text.

### Code-Beispiel

```python
def hello_world():
    print("Hello, World!")
```

## Tabelle

| Feature | Status |
|---------|--------|
| PDF Export | ✓ |
| DOCX Export | ✓ |
| Preview | ✓ |
"""
        self.text_input.insert("1.0", sample)
        self.update_preview()
    
    def on_css_preset_change(self):
        """Handle CSS preset change"""
        self.css_preset = self.css_preset_var.get()
        # Update converter's preset if it exists
        if self.converter:
            self.converter.css_preset = self.css_preset
        # Update preview with new style
        self.update_preview()
    
    def on_text_change(self, event=None):
        """Handle text changes with delayed preview update"""
        if hasattr(self, '_update_timer'):
            self.root.after_cancel(self._update_timer)
        self._update_timer = self.root.after(500, self.update_preview)
    
    def update_preview(self):
        """Update the HTML preview"""
        if not self.converter:
            self.status_var.set("Error: markdown library not installed")
            return
        
        try:
            md_content = self.text_input.get("1.0", tk.END).strip()
            if md_content:
                # Generate HTML from markdown
                try:
                    self.current_html = self.converter.markdown_to_html(md_content)
                except Exception as html_error:
                    print(f"[ERROR] HTML generation failed: {html_error}")
                    error_msg = f"Preview error: {str(html_error)}"
                    if self.use_html_preview and self.preview_html:
                        self.preview_html.load_html(f"<html><body><h1>Error</h1><p>{error_msg}</p></body></html>")
                    elif self.preview_text:
                        preview_content = md_content[:5000]  # Limit length
                        self.preview_text.config(state="normal")
                        self.preview_text.delete("1.0", tk.END)
                        self.preview_text.insert("1.0", preview_content)
                        self.preview_text.config(state="disabled")
                    self.status_var.set(error_msg)
                    return
                
                # Update preview based on available widget
                if self.use_html_preview and self.preview_html:
                    # Use embedded HTML browser
                    try:
                        self.preview_html.load_html(self.current_html)
                        self.status_var.set("Preview updated (HTML)")
                    except Exception as html_error:
                        print(f"[ERROR] Failed to load HTML in browser: {html_error}")
                        # Fallback to text preview if available
                        if self.preview_text:
                            preview_content = self._html_to_text_preview(self.current_html)
                            self.preview_text.config(state="normal")
                            self.preview_text.delete("1.0", tk.END)
                            self.preview_text.insert("1.0", preview_content)
                            self.preview_text.config(state="disabled")
                        self.status_var.set(f"Preview error: {str(html_error)}")
                else:
                    # Use text preview
                    try:
                        preview_content = self._html_to_text_preview(self.current_html)
                    except Exception as preview_error:
                        print(f"[ERROR] Preview conversion failed: {preview_error}")
                        # Fallback: show raw HTML (limited)
                        preview_content = self.current_html[:2000] + "\n\n[... HTML truncated ...]"
                    
                    # Ensure we have content to display
                    if not preview_content or preview_content.strip() == "":
                        # Last resort: show markdown directly
                        preview_content = md_content[:5000]
                        if len(md_content) > 5000:
                            preview_content += "\n\n[... content truncated ...]"
                    
                    if self.preview_text:
                        self.preview_text.config(state="normal")
                        self.preview_text.delete("1.0", tk.END)
                        self.preview_text.insert("1.0", preview_content)
                        
                        # Apply formatting (this should not modify text, only add tags)
                        try:
                            self._apply_preview_formatting()
                        except Exception as format_error:
                            # If formatting fails, at least show the text
                            print(f"[WARNING] Preview formatting error: {format_error}")
                        
                        self.preview_text.config(state="disabled")
                    
                    self.status_var.set("Preview updated")
            else:
                self.current_html = ""
                if self.use_html_preview and self.preview_html:
                    self.preview_html.load_html("<html><body><p>No content to preview</p></body></html>")
                elif self.preview_text:
                    self.preview_text.config(state="normal")
                    self.preview_text.delete("1.0", tk.END)
                    self.preview_text.config(state="disabled")
                self.status_var.set("No content to preview")
        except Exception as e:
            error_msg = f"Preview error: {str(e)}"
            self.status_var.set(error_msg)
            print(f"[ERROR] {error_msg}")
            import traceback
            traceback.print_exc()
            
            # Show error in preview
            if self.use_html_preview and self.preview_html:
                self.preview_html.load_html(f"<html><body><h1>Error</h1><p>Error generating preview:</p><p>{str(e)}</p><p>Please check console for details.</p></body></html>")
            elif self.preview_text:
                self.preview_text.config(state="normal")
                self.preview_text.delete("1.0", tk.END)
                self.preview_text.insert("1.0", f"Error generating preview:\n{str(e)}\n\nPlease check console for details.")
                self.preview_text.config(state="disabled")
    
    def _html_to_text_preview(self, html_content: str) -> str:
        """Convert HTML to readable text preview with formatting"""
        # Remove style and script tags completely
        html_content = re.sub(r'<style[^>]*>.*?</style>', '', html_content, flags=re.DOTALL | re.IGNORECASE)
        html_content = re.sub(r'<script[^>]*>.*?</script>', '', html_content, flags=re.DOTALL | re.IGNORECASE)
        
        # Remove head and meta tags
        html_content = re.sub(r'<head[^>]*>.*?</head>', '', html_content, flags=re.DOTALL | re.IGNORECASE)
        html_content = re.sub(r'<meta[^>]*>', '', html_content, flags=re.IGNORECASE)
        html_content = re.sub(r'<!DOCTYPE[^>]*>', '', html_content, flags=re.IGNORECASE)
        
        # Convert horizontal rules first (before headings)
        html_content = re.sub(r'<hr[^>]*>', '\n' + '-' * 50 + '\n', html_content, flags=re.IGNORECASE)
        
        # Convert headings with cleaner formatting
        html_content = re.sub(r'<h1[^>]*>(.*?)</h1>', r'\n\n\1\n' + '=' * 50 + '\n', html_content, flags=re.DOTALL | re.IGNORECASE)
        html_content = re.sub(r'<h2[^>]*>(.*?)</h2>', r'\n\n\1\n' + '-' * 50 + '\n', html_content, flags=re.DOTALL | re.IGNORECASE)
        html_content = re.sub(r'<h3[^>]*>(.*?)</h3>', r'\n\n\1\n' + '·' * 50 + '\n', html_content, flags=re.DOTALL | re.IGNORECASE)
        html_content = re.sub(r'<h4[^>]*>(.*?)</h4>', r'\n\1\n', html_content, flags=re.DOTALL | re.IGNORECASE)
        html_content = re.sub(r'<h[5-6][^>]*>(.*?)</h[5-6]>', r'\n\1\n', html_content, flags=re.DOTALL | re.IGNORECASE)
        
        # Convert code blocks
        html_content = re.sub(r'<pre[^>]*><code[^>]*>(.*?)</code></pre>', r'\n┌─ CODE BLOCK ─────────────────────┐\n\1\n└──────────────────────────────────┘\n', html_content, flags=re.DOTALL | re.IGNORECASE)
        html_content = re.sub(r'<pre[^>]*>(.*?)</pre>', r'\n┌─ CODE BLOCK ─────────────────────┐\n\1\n└──────────────────────────────────┘\n', html_content, flags=re.DOTALL | re.IGNORECASE)
        
        # Convert inline code
        html_content = re.sub(r'<code[^>]*>(.*?)</code>', r'`\1`', html_content, flags=re.DOTALL | re.IGNORECASE)
        
        # Convert lists - preserve nested structure
        html_content = re.sub(r'<ul[^>]*>', '\n', html_content, flags=re.IGNORECASE)
        html_content = re.sub(r'</ul>', '\n', html_content, flags=re.IGNORECASE)
        html_content = re.sub(r'<ol[^>]*>', '\n', html_content, flags=re.IGNORECASE)
        html_content = re.sub(r'</ol>', '\n', html_content, flags=re.IGNORECASE)
        # Handle list items - check for nested lists
        def replace_li(match):
            content = match.group(1)
            # Check if it contains nested lists
            if '<ul' in content or '<ol' in content:
                return f'  • {content}\n'
            return f'  • {content}\n'
        html_content = re.sub(r'<li[^>]*>(.*?)</li>', replace_li, html_content, flags=re.DOTALL | re.IGNORECASE)
        
        # Convert tables
        html_content = re.sub(r'<table[^>]*>', '\n┌─ TABLE ──────────────────────────┐\n', html_content, flags=re.IGNORECASE)
        html_content = re.sub(r'</table>', '└──────────────────────────────────┘\n', html_content, flags=re.IGNORECASE)
        html_content = re.sub(r'<tr[^>]*>', '│ ', html_content, flags=re.IGNORECASE)
        html_content = re.sub(r'</tr>', ' │\n', html_content, flags=re.IGNORECASE)
        html_content = re.sub(r'<th[^>]*>(.*?)</th>', r'[\1] ', html_content, flags=re.DOTALL | re.IGNORECASE)
        html_content = re.sub(r'<td[^>]*>(.*?)</td>', r'\1 | ', html_content, flags=re.DOTALL | re.IGNORECASE)
        
        # Convert blockquotes
        html_content = re.sub(r'<blockquote[^>]*>(.*?)</blockquote>', r'\n┌─ QUOTE ──────────────────────────┐\n\1\n└──────────────────────────────────┘\n', html_content, flags=re.DOTALL | re.IGNORECASE)
        
        # Convert paragraphs - but preserve spacing
        html_content = re.sub(r'<p[^>]*>(.*?)</p>', r'\1\n', html_content, flags=re.DOTALL | re.IGNORECASE)
        
        # Convert div containers (remove but preserve content)
        html_content = re.sub(r'<div[^>]*class="container"[^>]*>', '', html_content, flags=re.IGNORECASE)
        html_content = re.sub(r'<div[^>]*class="table-container"[^>]*>', '', html_content, flags=re.IGNORECASE)
        html_content = re.sub(r'<div[^>]*class="list-container"[^>]*>', '', html_content, flags=re.IGNORECASE)
        html_content = re.sub(r'</div>', '', html_content, flags=re.IGNORECASE)
        
        # Convert bold and italic - just extract text (formatting will be applied separately)
        html_content = re.sub(r'<strong[^>]*>(.*?)</strong>', r'\1', html_content, flags=re.DOTALL | re.IGNORECASE)
        html_content = re.sub(r'<b[^>]*>(.*?)</b>', r'\1', html_content, flags=re.DOTALL | re.IGNORECASE)
        html_content = re.sub(r'<em[^>]*>(.*?)</em>', r'\1', html_content, flags=re.DOTALL | re.IGNORECASE)
        html_content = re.sub(r'<i[^>]*>(.*?)</i>', r'\1', html_content, flags=re.DOTALL | re.IGNORECASE)
        
        # Convert links
        html_content = re.sub(r'<a[^>]*href="([^"]*)"[^>]*>(.*?)</a>', r'\2 (\1)', html_content, flags=re.DOTALL | re.IGNORECASE)
        html_content = re.sub(r'<a[^>]*>(.*?)</a>', r'\1', html_content, flags=re.DOTALL | re.IGNORECASE)
        
        # Convert citations
        html_content = re.sub(r'<span class="citation"[^>]*>(.*?)</span>', r'📄 \1', html_content, flags=re.DOTALL | re.IGNORECASE)
        
        # Remove remaining HTML tags
        html_content = re.sub(r'<[^>]+>', '', html_content)
        
        # Decode HTML entities
        html_content = html_content.replace('&amp;', '&')
        html_content = html_content.replace('&lt;', '<')
        html_content = html_content.replace('&gt;', '>')
        html_content = html_content.replace('&quot;', '"')
        html_content = html_content.replace('&#39;', "'")
        html_content = html_content.replace('&nbsp;', ' ')
        
        # Clean up whitespace - but preserve intentional spacing
        # Remove excessive blank lines (more than 2 consecutive)
        html_content = re.sub(r'\n\s*\n\s*\n+', '\n\n', html_content)
        # Remove leading/trailing whitespace from lines
        lines = html_content.split('\n')
        cleaned_lines = [line.rstrip() for line in lines]
        html_content = '\n'.join(cleaned_lines)
        html_content = html_content.strip()
        
        # Ensure we always return something
        if not html_content:
            return "Preview: Content is empty or could not be converted"
        
        return html_content
    
    def _apply_preview_formatting(self):
        """Apply text formatting tags to preview (read-only, doesn't modify text)"""
        try:
            content = self.preview_text.get("1.0", tk.END)
            if not content or content.strip() == "":
                return  # Nothing to format
            
            lines = content.split('\n')
            
            # Configure tags
            self.preview_text.tag_configure("heading1", font=("Segoe UI", 16, "bold"), foreground=UIColors.PRIMARY)
            self.preview_text.tag_configure("heading2", font=("Segoe UI", 13, "bold"), foreground=UIColors.PRIMARY)
            self.preview_text.tag_configure("heading3", font=("Segoe UI", 11, "bold"), foreground=UIColors.TEXT_PRIMARY)
            self.preview_text.tag_configure("heading4", font=("Segoe UI", 10, "bold"), foreground=UIColors.TEXT_SECONDARY)
            self.preview_text.tag_configure("separator", foreground=UIColors.BORDER)
            self.preview_text.tag_configure("code", font=("Consolas", 9), background=UIColors.BG_TERTIARY, foreground="#e74c3c")
            self.preview_text.tag_configure("bold", font=("Segoe UI", 10, "bold"))
            self.preview_text.tag_configure("italic", font=("Segoe UI", 10, "italic"))
            
            # Find and tag headings and separators
            for line_num, line in enumerate(lines, 1):
                if line_num > len(lines):
                    break
                    
                line_start = f"{line_num}.0"
                line_end = f"{line_num}.end"
                
                # Tag separators (lines with =, -, or ·)
                if line and all(c in '=-·' for c in line.strip()) and len(line.strip()) > 10:
                    try:
                        self.preview_text.tag_add("separator", line_start, line_end)
                        # Tag the line before as heading
                        if line_num > 1:
                            prev_line = f"{line_num-1}.0"
                            prev_line_end = f"{line_num-1}.end"
                            if '=' in line:
                                self.preview_text.tag_add("heading1", prev_line, prev_line_end)
                            elif '-' in line:
                                self.preview_text.tag_add("heading2", prev_line, prev_line_end)
                            elif '·' in line:
                                self.preview_text.tag_add("heading3", prev_line, prev_line_end)
                    except Exception:
                        pass  # Skip if line doesn't exist
                elif line.strip() and line_num > 1:
                    # Check if previous line was a separator
                    prev_line = lines[line_num - 2] if line_num > 1 else ""
                    if prev_line and all(c in '=-·' for c in prev_line.strip()) and len(prev_line.strip()) > 10:
                        try:
                            if '=' in prev_line:
                                self.preview_text.tag_add("heading1", line_start, line_end)
                            elif '-' in prev_line:
                                self.preview_text.tag_add("heading2", line_start, line_end)
                            elif '·' in prev_line:
                                self.preview_text.tag_add("heading3", line_start, line_end)
                        except Exception:
                            pass
                
                # Tag code blocks
                if '┌─ CODE' in line or '└─' in line or (line.strip().startswith('│') and 'CODE' in line):
                    try:
                        self.preview_text.tag_add("code", line_start, line_end)
                    except Exception:
                        pass
                
                # Tag inline code (backticks) - process from end to start
                if '`' in line and not line.startswith('┌─'):
                    try:
                        line_content = self.preview_text.get(line_start, line_end)
                        code_matches = list(re.finditer(r'`([^`]+)`', line_content))
                        for match in reversed(code_matches):
                            code_text = match.group(1)
                            start_pos = f"{line_num}.{match.start()}"
                            end_pos = f"{line_num}.{match.end()}"
                            # Replace backticks with text
                            self.preview_text.delete(start_pos, end_pos)
                            self.preview_text.insert(start_pos, code_text)
                            # Apply code formatting
                            new_end = f"{line_num}.{match.start() + len(code_text)}"
                            self.preview_text.tag_add("code", start_pos, new_end)
                    except Exception:
                        pass  # Skip if there's an error
        except Exception as e:
            # Don't break preview if formatting fails
            print(f"[WARNING] Preview formatting error: {e}")
    
    def open_file(self):
        """Open and load a markdown file"""
        file_path = filedialog.askopenfilename(
            title="Select Markdown File",
            filetypes=[("Markdown files", "*.md *.markdown"), ("All files", "*.*")]
        )
        
        if file_path:
            self.load_file(file_path)
    
    def load_file(self, file_path):
        """Load content from a file (Feature 1: Encoding Detection)"""
        try:
            # Use encoding detection instead of hardcoded UTF-8
            content = _read_file_with_encoding_detection(file_path)
            
            self.text_input.delete("1.0", tk.END)
            self.text_input.insert("1.0", content)
            self.update_preview()
            
            # Track file path for save dialog
            self.current_file_path = file_path
            
            filename = os.path.basename(file_path)
            self.status_var.set(f"Loaded: {filename}")
            
        except Exception as e:
            messagebox.showerror("Error", f"Failed to load file: {str(e)}")
    
    def save_pdf(self):
        """Save content as PDF"""
        if not self.converter:
            messagebox.showerror("Error", "Markdown library not installed. Please install with: pip install markdown")
            return
        
        md_content = self.text_input.get("1.0", tk.END).strip()
        if not md_content:
            messagebox.showwarning("Warning", "No content to export!")
            return
        
        # Determine initial directory and filename from current file or use current working directory
        initial_dir = None
        initial_file = None
        if self.current_file_path:
            initial_dir = os.path.dirname(os.path.abspath(self.current_file_path))
            # Auto-set filename based on loaded file basename
            base_name = os.path.splitext(os.path.basename(self.current_file_path))[0]
            initial_file = f"{base_name}.pdf"
        elif os.path.exists(os.getcwd()):
            initial_dir = os.getcwd()
        
        file_path = filedialog.asksaveasfilename(
            title="Save as PDF",
            defaultextension=".pdf",
            filetypes=[("PDF files", "*.pdf"), ("All files", "*.*")],
            initialdir=initial_dir,
            initialfile=initial_file
        )
        
        if file_path:
            self.status_var.set("Exporting to PDF...")
            # Set wait cursor during processing
            self.root.config(cursor="wait")
            self.root.update()
            
            try:
                orientation = self.pdf_orientation.get()
                engine = self.pdf_engine.get()
                
                success = False
                # Try WeasyPrint first if selected (Feature 6)
                if engine == "weasyprint" and _try_import_weasyprint():
                    success = self.converter.markdown_to_pdf_weasyprint(md_content, file_path, orientation=orientation)
                elif engine == "browser":
                    html = self.current_html or self.converter.markdown_to_html(md_content)
                    success = self.converter.html_to_pdf_browser(html, file_path, orientation=orientation)
                
                # Fallback to ReportLab if other methods failed
                if not success and REPORTLAB_AVAILABLE:
                    success = self.converter.markdown_to_pdf_reportlab(md_content, file_path, orientation, keep_icons=True)
                
                if success:
                    self.status_var.set(f"PDF saved: {os.path.basename(file_path)}")
                    # Auto-open PDF if enabled
                    if self.settings_manager.get('autoopen_enabled', True):
                        try:
                            os.startfile(file_path)  # Windows
                        except AttributeError:
                            # Unix/Linux/Mac
                            try:
                                subprocess.run(['xdg-open', file_path], check=False)
                            except Exception:
                                try:
                                    subprocess.run(['open', file_path], check=False)  # Mac
                                except Exception:
                                    pass
                    else:
                        messagebox.showinfo("Success", f"PDF exported successfully!\n{file_path}")
                else:
                    self.status_var.set("PDF export failed")
                    messagebox.showerror("Error", "PDF export failed. Please check console for details.")
            finally:
                # Always restore cursor
                self.root.config(cursor="")
    
    def save_docx(self):
        """Save content as DOCX"""
        if not self.converter:
            messagebox.showerror("Error", "Markdown library not installed. Please install with: pip install markdown")
            return
        
        md_content = self.text_input.get("1.0", tk.END).strip()
        if not md_content:
            messagebox.showwarning("Warning", "No content to export!")
            return
        
        # Determine initial directory and filename from current file or use current working directory
        initial_dir = None
        initial_file = None
        if self.current_file_path:
            initial_dir = os.path.dirname(os.path.abspath(self.current_file_path))
            # Auto-set filename based on loaded file basename
            base_name = os.path.splitext(os.path.basename(self.current_file_path))[0]
            initial_file = f"{base_name}.docx"
        elif os.path.exists(os.getcwd()):
            initial_dir = os.getcwd()
        
        file_path = filedialog.asksaveasfilename(
            title="Save as DOCX",
            defaultextension=".docx",
            filetypes=[("Word documents", "*.docx"), ("All files", "*.*")],
            initialdir=initial_dir,
            initialfile=initial_file
        )
        
        if file_path:
            self.status_var.set("Exporting to DOCX...")
            # Set wait cursor during processing
            self.root.config(cursor="wait")
            self.root.update()
            
            try:
                if self.converter.markdown_to_docx(md_content, file_path):
                    self.status_var.set(f"DOCX saved: {os.path.basename(file_path)}")
                    # Auto-open DOCX if enabled
                    if self.settings_manager.get('autoopen_enabled', True):
                        try:
                            os.startfile(file_path)  # Windows
                        except AttributeError:
                            # Unix/Linux/Mac
                            try:
                                subprocess.run(['xdg-open', file_path], check=False)
                            except Exception:
                                try:
                                    subprocess.run(['open', file_path], check=False)  # Mac
                                except Exception:
                                    pass
                    else:
                        messagebox.showinfo("Success", f"DOCX exported successfully!\n{file_path}")
                else:
                    self.status_var.set("DOCX export failed")
            finally:
                # Always restore cursor
                self.root.config(cursor="")
    
    def show_settings_dialog(self):
        """Show settings dialog for optional features"""
        dialog = tk.Toplevel(self.root)
        dialog.title("Settings")
        dialog.geometry("600x500")
        dialog.configure(bg=UIColors.BG_SECONDARY)
        dialog.transient(self.root)
        dialog.grab_set()
        dialog.lift()
        dialog.focus_force()
        
        # Center on parent
        dialog.update_idletasks()
        x = self.root.winfo_x() + (self.root.winfo_width() - 600) // 2
        y = self.root.winfo_y() + (self.root.winfo_height() - 500) // 2
        dialog.geometry(f"+{x}+{y}")
        
        # Main frame
        main_frame = tk.Frame(dialog, bg=UIColors.BG_SECONDARY, padx=UISpacing.MD, pady=UISpacing.MD)
        main_frame.pack(fill='both', expand=True)
        
        # Title
        title_label = tk.Label(
            main_frame,
            text="⚙️ Converter Settings",
            font=UIFonts.TITLE,
            bg=UIColors.BG_SECONDARY,
            fg=UIColors.PRIMARY
        )
        title_label.pack(pady=(0, UISpacing.LG))
        
        # Settings variables
        page_breaks_var = tk.BooleanVar(value=self.settings_manager.get('page_breaks_enabled', True))
        preprocessing_var = tk.BooleanVar(value=self.settings_manager.get('preprocessing_enabled', True))
        advanced_docx_var = tk.BooleanVar(value=self.settings_manager.get('advanced_docx_enabled', True))
        font_size_var = tk.BooleanVar(value=self.settings_manager.get('font_size_options_enabled', True))
        font_size_value = tk.IntVar(value=self.settings_manager.get('default_font_size', 10))
        narrow_margins_var = tk.BooleanVar(value=self.settings_manager.get('docx_narrow_margins', False))
        auto_h2_var = tk.BooleanVar(value=self.settings_manager.get('auto_page_break_h2', True))
        auto_h3_var = tk.BooleanVar(value=self.settings_manager.get('auto_page_break_h3', True))
        autoopen_var = tk.BooleanVar(value=self.settings_manager.get('autoopen_enabled', True))
        
        # Feature 2: Page Breaks
        page_breaks_frame = tk.LabelFrame(
            main_frame,
            text=" Page Breaks (Feature 2) ",
            font=UIFonts.HEADING,
            bg=UIColors.BG_PRIMARY,
            fg=UIColors.TEXT_PRIMARY,
            padx=UISpacing.MD,
            pady=UISpacing.SM
        )
        page_breaks_frame.pack(fill='x', pady=UISpacing.SM)
        
        tk.Checkbutton(
            page_breaks_frame,
            text="Enable page break support",
            variable=page_breaks_var,
            bg=UIColors.BG_PRIMARY,
            font=UIFonts.BODY
        ).pack(anchor='w', pady=UISpacing.XS)
        tk.Checkbutton(
            page_breaks_frame,
            text="Auto page break before numbered H2 headings (e.g., '1.', '2.')",
            variable=auto_h2_var,
            bg=UIColors.BG_PRIMARY,
            font=UIFonts.BODY
        ).pack(anchor='w', pady=UISpacing.XS)
        tk.Checkbutton(
            page_breaks_frame,
            text="Auto page break before numbered H3 sub-sections (e.g., '1.1', '2.3')",
            variable=auto_h3_var,
            bg=UIColors.BG_PRIMARY,
            font=UIFonts.BODY
        ).pack(anchor='w', pady=UISpacing.XS)
        
        # Feature 3: Preprocessing
        preprocessing_frame = tk.LabelFrame(
            main_frame,
            text=" Content Preprocessing (Feature 3) ",
            font=UIFonts.HEADING,
            bg=UIColors.BG_PRIMARY,
            fg=UIColors.TEXT_PRIMARY,
            padx=UISpacing.MD,
            pady=UISpacing.SM
        )
        preprocessing_frame.pack(fill='x', pady=UISpacing.SM)
        
        tk.Checkbutton(
            preprocessing_frame,
            text="Remove horizontal rules (---, ***) and normalize line breaks",
            variable=preprocessing_var,
            bg=UIColors.BG_PRIMARY,
            font=UIFonts.BODY
        ).pack(anchor='w', pady=UISpacing.XS)
        
        # Feature 4: Advanced DOCX
        docx_frame = tk.LabelFrame(
            main_frame,
            text=" Advanced DOCX Features (Feature 4) ",
            font=UIFonts.HEADING,
            bg=UIColors.BG_PRIMARY,
            fg=UIColors.TEXT_PRIMARY,
            padx=UISpacing.MD,
            pady=UISpacing.SM
        )
        docx_frame.pack(fill='x', pady=UISpacing.SM)
        
        tk.Checkbutton(
            docx_frame,
            text="Enable advanced DOCX features (bookmarks, hyperlinks, formatting)",
            variable=advanced_docx_var,
            bg=UIColors.BG_PRIMARY,
            font=UIFonts.BODY
        ).pack(anchor='w', pady=UISpacing.XS)
        tk.Checkbutton(
            docx_frame,
            text="Use narrow margins (0.5 inch) in DOCX",
            variable=narrow_margins_var,
            bg=UIColors.BG_PRIMARY,
            font=UIFonts.BODY
        ).pack(anchor='w', pady=UISpacing.XS)
        
        # Feature 5: Font Size
        font_frame = tk.LabelFrame(
            main_frame,
            text=" Font Size Options (Feature 5) ",
            font=UIFonts.HEADING,
            bg=UIColors.BG_PRIMARY,
            fg=UIColors.TEXT_PRIMARY,
            padx=UISpacing.MD,
            pady=UISpacing.SM
        )
        font_frame.pack(fill='x', pady=UISpacing.SM)
        
        tk.Checkbutton(
            font_frame,
            text="Enable font size options",
            variable=font_size_var,
            bg=UIColors.BG_PRIMARY,
            font=UIFonts.BODY
        ).pack(anchor='w', pady=UISpacing.XS)
        
        font_size_frame = tk.Frame(font_frame, bg=UIColors.BG_PRIMARY)
        font_size_frame.pack(anchor='w', padx=UISpacing.LG, pady=UISpacing.XS)
        tk.Label(font_size_frame, text="Default font size:", bg=UIColors.BG_PRIMARY, font=UIFonts.BODY).pack(side='left', padx=UISpacing.SM)
        for size in [9, 10, 11, 12, 14]:
            tk.Radiobutton(
                font_size_frame,
                text=f"{size}pt",
                variable=font_size_value,
                value=size,
                bg=UIColors.BG_PRIMARY,
                font=UIFonts.BODY
            ).pack(side='left', padx=UISpacing.XS)
        
        # Auto-open settings
        autoopen_frame = tk.LabelFrame(
            main_frame,
            text=" Export Options ",
            font=UIFonts.HEADING,
            bg=UIColors.BG_PRIMARY,
            fg=UIColors.TEXT_PRIMARY,
            padx=UISpacing.MD,
            pady=UISpacing.SM
        )
        autoopen_frame.pack(fill='x', pady=UISpacing.SM)
        
        tk.Checkbutton(
            autoopen_frame,
            text="Automatically open PDF/DOCX after generation",
            variable=autoopen_var,
            bg=UIColors.BG_PRIMARY,
            font=UIFonts.BODY
        ).pack(anchor='w', pady=UISpacing.XS)
        
        # Buttons
        button_frame = tk.Frame(main_frame, bg=UIColors.BG_SECONDARY)
        button_frame.pack(fill='x', pady=(UISpacing.LG, 0))
        
        def save_settings():
            self.settings_manager.set('page_breaks_enabled', page_breaks_var.get())
            self.settings_manager.set('preprocessing_enabled', preprocessing_var.get())
            self.settings_manager.set('advanced_docx_enabled', advanced_docx_var.get())
            self.settings_manager.set('font_size_options_enabled', font_size_var.get())
            self.settings_manager.set('default_font_size', font_size_value.get())
            self.settings_manager.set('docx_narrow_margins', narrow_margins_var.get())
            self.settings_manager.set('auto_page_break_h2', auto_h2_var.get())
            self.settings_manager.set('auto_page_break_h3', auto_h3_var.get())
            self.settings_manager.set('autoopen_enabled', autoopen_var.get())
            
            # Update converter settings
            if self.converter:
                self.converter.settings = self.settings_manager
            
            self.settings_manager.save_settings()
            dialog.destroy()
            self.status_var.set("Settings saved")
        
        self.create_rounded_button(button_frame, "Save", save_settings, "success").pack(side='right', padx=UISpacing.SM)
        self.create_rounded_button(button_frame, "Cancel", dialog.destroy, "secondary").pack(side='right', padx=UISpacing.SM)
    
    def on_closing(self):
        """Handle window closing - save settings"""
        self.settings_manager.save_settings()
        self.root.destroy()
    
    def open_in_browser(self):
        """Open current HTML in browser"""
        if not self.current_html:
            messagebox.showwarning("Warning", "No content to preview!")
            return
        
        try:
            if not self.browser_preview_path:
                temp_dir = tempfile.gettempdir()
                self.browser_preview_path = os.path.join(temp_dir, "md_converter_live_preview.html")
                
                def _cleanup():
                    try:
                        if self.browser_preview_path and os.path.exists(self.browser_preview_path):
                            os.unlink(self.browser_preview_path)
                    except Exception:
                        pass
                atexit.register(_cleanup)
            
            with open(self.browser_preview_path, "w", encoding="utf-8") as f:
                f.write(self.current_html)
            
            file_url = Path(self.browser_preview_path).absolute().as_uri()
            webbrowser.open(file_url)
            self.status_var.set("Browser preview opened")
            
        except Exception as e:
            messagebox.showerror("Error", f"Failed to open in browser: {str(e)}")
    
    def run(self):
        """Start the GUI main loop"""
        self.root.mainloop()


def install_missing_dependencies():
    """Automatically install missing dependencies"""
    missing_deps = []
    
    if not MARKDOWN_AVAILABLE:
        missing_deps.append("markdown>=3.4.0")
    if not TKINTERWEB_AVAILABLE:
        missing_deps.append("tkinterweb>=3.0.0")
    if not DOCX_AVAILABLE:
        missing_deps.append("python-docx>=1.1.0")
    if not REPORTLAB_AVAILABLE:
        missing_deps.append("reportlab>=4.0.0")
    
    if missing_deps:
        print(f"[INFO] Missing dependencies detected: {', '.join(missing_deps)}")
        print("[INFO] Attempting to install missing dependencies...")
        
        try:
            result = subprocess.run(
                [sys.executable, "-m", "pip", "install", "--quiet", "--upgrade"] + missing_deps,
                capture_output=True,
                text=True,
                timeout=120
            )
            
            if result.returncode == 0:
                print("[SUCCESS] Missing dependencies installed successfully.")
                print("[INFO] Please restart the application for changes to take effect.")
                return True
            else:
                print(f"[WARNING] Failed to install some dependencies: {result.stderr}")
                print(f"[INFO] Please install manually with: pip install {' '.join(missing_deps)}")
                return False
        except Exception as e:
            print(f"[ERROR] Failed to install dependencies: {e}")
            print(f"[INFO] Please install manually with: pip install {' '.join(missing_deps)}")
            return False
    
    return True

def main():
    """Main entry point"""
    # Try to install missing dependencies automatically
    install_missing_dependencies()
    
    app = MarkdownConverterGUI()
    app.run()


if __name__ == "__main__":
    main()
